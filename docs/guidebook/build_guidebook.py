"""Perakit Buku Panduan (DOCX) aplikasi Akuntansi Tani Padi.

Skrip ini menyusun dokumen Word terstruktur (sampul, kata pengantar, daftar
isi, BAB I-III) dan menyematkan gambar kode (carbon) serta tangkapan layar
aplikasi nyata. Idempoten: menulis ulang berkas keluaran setiap dijalankan.
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).resolve().parent
CODE = BASE / "assets" / "code"
SHOT = BASE / "assets" / "screenshots"
FLOW = BASE / "assets" / "flowchart.png"
OUT = BASE / "Buku-Panduan-Akuntansi-Tani-Padi.docx"

GREEN = RGBColor(0x1B, 0x5E, 0x20)
GREEN2 = RGBColor(0x2E, 0x7D, 0x32)
DARK = RGBColor(0x21, 0x21, 0x21)
GREY = RGBColor(0x55, 0x55, 0x55)

BODY_FONT = "Calibri"
BODY_SIZE = Pt(11.5)

_fig = {"n": 0}


# --------------------------------------------------------------------------- helpers
def set_cell_no_border(doc):
    pass


def _shade(el, color_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    el.append(shd)


def style_base(doc):
    st = doc.styles["Normal"]
    st.font.name = BODY_FONT
    st.font.size = BODY_SIZE
    st.font.color.rgb = DARK
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.4
    pf.space_after = Pt(6)

    for name, size, color in (("Heading 1", 18, GREEN), ("Heading 2", 14, GREEN2),
                              ("Heading 3", 12.5, GREEN2)):
        s = doc.styles[name]
        s.font.name = BODY_FONT
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(14)
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.keep_with_next = True


def set_a4_margins(section):
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)


def footer_with_pagenum(section, label):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label + "  |  Halaman ")
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    r.font.name = BODY_FONT
    add_page_field(p)


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    return p


def h2(doc, text):
    return doc.add_heading(text, level=2)


def h3(doc, text):
    return doc.add_heading(text, level=3)


def para(doc, text, justify=True, italic=False, size=None, color=None, bold=False):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.name = BODY_FONT
    if size:
        r.font.size = size
    if color:
        r.font.color.rgb = color
    return p


def bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_lead:
        rb = p.add_run(bold_lead)
        rb.bold = True
        rb.font.name = BODY_FONT
    r = p.add_run(text)
    r.font.name = BODY_FONT
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text).font.name = BODY_FONT
    return p


def figure(doc, path, caption, width_cm=15.5):
    _fig["n"] += 1
    n = _fig["n"]
    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.space_before = Pt(8)
    pic_p.paragraph_format.space_after = Pt(2)
    run = pic_p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    rc = cap.add_run(f"Gambar {n}. {caption}")
    rc.italic = True
    rc.font.size = Pt(10)
    rc.font.color.rgb = GREY
    rc.font.name = BODY_FONT
    return n


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r' TOC \o "1-3" \h \z \u '
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "Klik kanan lalu pilih \u201cUpdate Field\u201d untuk memperbarui daftar isi."
    fld3 = OxmlElement("w:fldChar")
    fld3.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)
    run._r.append(t)
    run._r.append(fld3)


def enable_update_fields(doc):
    settings = doc.settings.element
    el = OxmlElement("w:updateFields")
    el.set(qn("w:val"), "true")
    settings.append(el)


def kv_table(doc, rows, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        rp = hdr[i].paragraphs[0].add_run(htext)
        rp.bold = True
        rp.font.size = Pt(10)
        rp.font.name = BODY_FONT
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            rr = cells[i].paragraphs[0].add_run(str(val))
            rr.font.size = Pt(10)
            rr.font.name = BODY_FONT
    return table


def page_break(doc):
    doc.add_page_break()


# --------------------------------------------------------------------------- build
doc = Document()
style_base(doc)

# ===== SAMPUL =====
sec0 = doc.sections[0]
set_a4_margins(sec0)

doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("PENERAPAN SISTEM INFORMASI AKUNTANSI\nUNTUK PENCATATAN KEUANGAN USAHA TANI PADI")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = GREEN
r.font.name = BODY_FONT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("(BUKU PANDUAN)")
rs.bold = True
rs.font.size = Pt(16)
rs.font.color.rgb = GREEN2
rs.font.name = BODY_FONT

doc.add_paragraph()
logo_p = doc.add_paragraph()
logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
logo_p.add_run().add_picture(str(FLOW), width=Cm(6.5))
cap = doc.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
rc = cap.add_run("Aplikasi Akuntansi Tani Padi berbasis Streamlit")
rc.italic = True
rc.font.size = Pt(10)
rc.font.color.rgb = GREY
rc.font.name = BODY_FONT

doc.add_paragraph()
auth = doc.add_paragraph()
auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
ra = auth.add_run("Disusun oleh:\nTim Pengembang Aplikasi")
ra.bold = True
ra.font.size = Pt(13)
ra.font.name = BODY_FONT

yr = doc.add_paragraph()
yr.alignment = WD_ALIGN_PARAGRAPH.CENTER
ry = yr.add_run("2026")
ry.bold = True
ry.font.size = Pt(13)
ry.font.color.rgb = GREEN
ry.font.name = BODY_FONT

page_break(doc)

# ===== KATA PENGANTAR =====
h1(doc, "KATA PENGANTAR")
para(doc, "Puji syukur kami panjatkan ke hadirat Tuhan Yang Maha Esa atas "
          "limpahan rahmat dan karunia-Nya sehingga Buku Panduan Aplikasi Akuntansi "
          "Tani Padi ini dapat diselesaikan dengan baik. Buku panduan ini disusun "
          "sebagai dokumentasi resmi yang menjelaskan konsep, perancangan, hingga cara "
          "penggunaan aplikasi pencatatan keuangan untuk usaha tani padi.")
para(doc, "Aplikasi Akuntansi Tani Padi dikembangkan untuk membantu petani dan "
          "pelaku usaha tani dalam mencatat transaksi keuangan secara rapi dan "
          "sistematis, mulai dari pencatatan jurnal umum hingga penyusunan laporan "
          "keuangan lengkap. Dengan antarmuka yang sederhana berbasis web, aplikasi "
          "ini diharapkan dapat menjembatani kebutuhan pencatatan keuangan yang akurat "
          "tanpa memerlukan latar belakang akuntansi yang mendalam.")
para(doc, "Buku panduan ini memuat penjelasan menyeluruh, mulai dari latar belakang, "
          "arsitektur aplikasi, alur kerja sistem, struktur basis data, implementasi "
          "kode program, hingga langkah penggunaan setiap fitur yang dilengkapi dengan "
          "tangkapan layar aplikasi. Penjelasan kode disertai potongan program yang "
          "sesungguhnya agar pembaca dapat memahami logika di balik setiap fungsi.")
para(doc, "Kami menyadari bahwa buku panduan ini masih memiliki kekurangan. Oleh "
          "karena itu, kritik dan saran yang membangun sangat kami harapkan demi "
          "penyempurnaan di masa mendatang. Semoga buku panduan ini bermanfaat bagi "
          "semua pihak yang menggunakannya.")
para(doc, "Tim Pengembang Aplikasi", justify=False, bold=True)

page_break(doc)

# ===== DAFTAR ISI =====
h1(doc, "DAFTAR ISI")
add_toc(doc)
page_break(doc)

# ===== BAB I PENDAHULUAN =====
h1(doc, "BAB I  PENDAHULUAN")

h2(doc, "1.1  Latar Belakang")
para(doc, "Usaha tani padi merupakan salah satu sektor penting dalam perekonomian, "
          "namun pencatatan keuangannya sering kali masih dilakukan secara manual atau "
          "bahkan tidak dicatat sama sekali. Akibatnya, petani kesulitan mengetahui "
          "secara pasti besarnya pendapatan, beban, serta laba yang diperoleh dari "
          "setiap musim tanam. Ketidakteraturan pencatatan ini menyulitkan pengambilan "
          "keputusan dan evaluasi usaha.")
para(doc, "Sistem informasi akuntansi hadir sebagai solusi untuk mengelola data "
          "keuangan secara terstruktur. Dengan memanfaatkan teknologi, proses "
          "pencatatan transaksi, penggolongan akun, hingga penyusunan laporan keuangan "
          "dapat dilakukan secara otomatis dan akurat. Aplikasi Akuntansi Tani Padi "
          "dibangun untuk menjawab kebutuhan tersebut dengan menyediakan siklus "
          "akuntansi lengkap dalam satu aplikasi berbasis web yang ringan dan mudah "
          "digunakan.")

h2(doc, "1.2  Tujuan")
numbered(doc, "Menyediakan sarana pencatatan transaksi keuangan usaha tani padi yang "
              "rapi, akurat, dan mudah digunakan.")
numbered(doc, "Menerapkan siklus akuntansi lengkap, mulai dari jurnal umum hingga "
              "laporan keuangan, secara otomatis.")
numbered(doc, "Membantu pelaku usaha tani memahami posisi keuangan usahanya melalui "
              "laporan laba rugi, perubahan ekuitas, posisi keuangan, dan arus kas.")
numbered(doc, "Menyediakan dokumentasi teknis yang menjelaskan arsitektur, kode "
              "program, dan cara penggunaan aplikasi secara menyeluruh.")

h2(doc, "1.3  Manfaat")
bullet(doc, "Pencatatan keuangan menjadi lebih tertib dan dapat dipertanggungjawabkan.",
       bold_lead="Bagi pengguna: ")
bullet(doc, "Laporan keuangan dihasilkan secara otomatis sehingga menghemat waktu dan "
            "mengurangi kesalahan hitung.", bold_lead="Bagi usaha tani: ")
bullet(doc, "Menjadi acuan pengembangan dan pemeliharaan aplikasi pada masa mendatang.",
       bold_lead="Bagi pengembang: ")

h2(doc, "1.4  Ruang Lingkup")
para(doc, "Buku panduan ini mencakup penjelasan mengenai teknologi yang digunakan, "
          "alur kerja aplikasi, struktur basis data, implementasi kode program inti, "
          "serta panduan penggunaan setiap menu. Pembahasan difokuskan pada aplikasi "
          "Akuntansi Tani Padi yang dibangun menggunakan bahasa pemrograman Python "
          "dengan kerangka kerja Streamlit dan basis data SQLite.")

page_break(doc)

# ===== BAB II ISI =====
h1(doc, "BAB II  PEMBAHASAN")

# ---- 2.1 ----
h2(doc, "2.1  Gambaran Umum Aplikasi")
para(doc, "Aplikasi Akuntansi Tani Padi adalah sistem informasi akuntansi berbasis web "
          "yang menerapkan siklus akuntansi secara menyeluruh untuk usaha tani padi. "
          "Aplikasi dijalankan melalui peramban (browser) dan menampilkan data dalam "
          "bentuk tabel serta laporan keuangan yang mudah dibaca. Seluruh data "
          "transaksi disimpan secara permanen pada basis data lokal.")

h3(doc, "2.1.1  Teknologi dan Pustaka yang Digunakan")
para(doc, "Aplikasi ini dibangun menggunakan beberapa teknologi dan pustaka utama "
          "sebagai berikut:")
kv_table(doc, [
    ["Python 3", "Bahasa pemrograman utama yang digunakan untuk seluruh logika aplikasi."],
    ["Streamlit", "Kerangka kerja antarmuka web untuk membangun halaman aplikasi secara cepat menggunakan Python murni."],
    ["SQLite", "Basis data lokal berbasis berkas (tani_padi.db) untuk menyimpan data pengguna dan jurnal."],
    ["Pandas", "Pustaka pengolahan data tabular yang digunakan untuk menyusun dan menampilkan tabel laporan."],
    ["Decimal", "Modul bawaan Python untuk perhitungan uang dengan presisi tinggi sehingga tidak terjadi galat pembulatan."],
    ["hashlib", "Modul bawaan Python untuk mengamankan kata sandi pengguna dengan algoritme PBKDF2-HMAC-SHA256."],
], headers=["Teknologi / Pustaka", "Keterangan"])

h3(doc, "2.1.2  Fitur Utama")
bullet(doc, "Autentikasi pengguna (login dan pendaftaran) dengan kata sandi terenkripsi.")
bullet(doc, "Pencatatan transaksi dengan sistem berpasangan (debit dan kredit) yang divalidasi otomatis.")
bullet(doc, "Penyajian siklus akuntansi lengkap dalam 11 tahap laporan.")
bullet(doc, "Dasbor ringkasan keuangan (pendapatan, beban, laba bersih, dan kas).")
bullet(doc, "Fitur ubah dan hapus transaksi yang fleksibel.")

h3(doc, "2.1.3  Siklus Akuntansi 11 Tahap")
para(doc, "Aplikasi menyajikan keseluruhan siklus akuntansi yang terdiri atas sebelas "
          "tahap laporan berikut:")
for i, nm in enumerate([
    "Jurnal Umum", "Buku Besar", "Neraca Saldo", "Jurnal Penyesuaian",
    "Neraca Saldo Setelah Penyesuaian", "Laporan Laba Rugi",
    "Laporan Perubahan Ekuitas", "Laporan Posisi Keuangan (Neraca)",
    "Laporan Arus Kas", "Jurnal Penutup", "Neraca Saldo Setelah Penutupan",
], start=1):
    numbered(doc, nm)

# ---- 2.2 ----
h2(doc, "2.2  Alur Kerja Aplikasi (Flowchart)")
para(doc, "Diagram alur berikut menggambarkan jalannya aplikasi sejak dijalankan "
          "hingga pengguna dapat mengakses laporan dan mengelola transaksi. Ketika "
          "aplikasi dijalankan, konfigurasi halaman dan basis data disiapkan terlebih "
          "dahulu. Pengguna kemudian diarahkan ke halaman login. Setelah berhasil "
          "masuk, pengguna dapat memilih menu untuk melihat laporan atau mengelola "
          "transaksi.")
figure(doc, FLOW, "Diagram alur kerja aplikasi Akuntansi Tani Padi.", width_cm=13.5)
para(doc, "Inti dari alur ini adalah pemeriksaan status login dan validasi transaksi. "
          "Setiap transaksi yang diinput wajib memenuhi syarat keseimbangan, yaitu "
          "total debit harus sama dengan total kredit, sebelum data disimpan ke basis "
          "data.")

# ---- 2.3 ----
h2(doc, "2.3  Persiapan Lingkungan di Windows")
para(doc, "Bagian ini menjelaskan cara menyiapkan dan menjalankan aplikasi pada "
          "komputer ber-sistem operasi Windows, langkah demi langkah, mulai dari "
          "memasang Python sampai aplikasi terbuka di peramban. Urutan langkahnya "
          "adalah: (1) memasang Python dan mencentang \u201cAdd Python to PATH\u201d, "
          "(2) membuat virtual environment, (3) mengaktifkan virtual environment, "
          "(4) memasang pustaka yang dibutuhkan, dan (5) menjalankan aplikasi. "
          "Ikuti langkah ini secara berurutan agar tidak ada tahap yang terlewat. "
          "Semua perintah pada bagian ini dijalankan melalui Command Prompt (cmd) "
          "atau PowerShell bawaan Windows.")

h3(doc, "2.3.1  Memasang Python di Windows")
para(doc, "Python adalah bahasa pemrograman yang menjalankan seluruh aplikasi, "
          "sehingga harus dipasang terlebih dahulu. Unduh pemasang (installer) "
          "Python untuk Windows dari situs resmi di alamat "
          "https://www.python.org/downloads/ , lalu jalankan berkas pemasang yang "
          "telah diunduh. Pada jendela pertama pemasang, WAJIB mencentang kotak "
          "\u201cAdd Python to PATH\u201d (Tambahkan Python ke PATH) terlebih dahulu "
          "sebelum menekan tombol \u201cInstall Now\u201d. Centang ini membuat "
          "perintah python dan py dapat dikenali dari Command Prompt mana pun. "
          "Setelah pemasangan selesai, buka Command Prompt (ketik \u201ccmd\u201d "
          "pada menu Start), lalu ketik perintah verifikasi berikut untuk "
          "memastikan Python sudah terpasang dengan benar.")
figure(doc, CODE / "Command_Prompt_memeriksa_Python_Windows.png",
       "Memverifikasi pemasangan Python di Command Prompt Windows.")
para(doc, "Apabila perintah py --version atau python --version menampilkan nomor "
          "versi (misalnya Python 3.12.x), berarti Python sudah berhasil dipasang "
          "dan siap digunakan. Bila muncul pesan bahwa perintah tidak dikenali, "
          "ulangi pemasangan dan pastikan kotak \u201cAdd Python to PATH\u201d "
          "benar-benar tercentang.")

h3(doc, "2.3.2  Membuat dan Mengaktifkan Virtual Environment di Windows")
para(doc, "Virtual environment adalah folder khusus yang menampung pustaka aplikasi "
          "secara terpisah, sehingga pustaka yang dipasang tidak bercampur atau "
          "mengganggu instalasi Python lain di komputer. Pada Windows, virtual "
          "environment dibuat menggunakan Python Launcher bawaan, yaitu perintah "
          "py, dengan menjalankan py -m venv venv (kata venv yang terakhir adalah "
          "nama folder yang akan dibuat). Setelah folder terbentuk, virtual "
          "environment perlu diaktifkan agar perintah pip dan streamlit dijalankan "
          "di dalam lingkungan tersebut. Cara mengaktifkannya berbeda sedikit "
          "tergantung jenis terminal yang dipakai, seperti terlihat pada gambar "
          "berikut.")
figure(doc, CODE / "Command_Prompt_membuat_dan_mengaktifkan_virtual_environment_Windows.png",
       "Membuat dan mengaktifkan virtual environment di Windows.")
bullet(doc, "jalankan venv\\Scripts\\activate.bat (atau cukup venv\\Scripts\\activate).",
       bold_lead="Pada Command Prompt (cmd): ")
bullet(doc, "jalankan .\\venv\\Scripts\\Activate.ps1. Bila PowerShell menolak karena "
            "kebijakan keamanan, jalankan dulu perintah "
            "Set-ExecutionPolicy -Scope CurrentUser RemoteSigned lalu coba lagi.",
       bold_lead="Pada PowerShell: ")
bullet(doc, "virtual environment berhasil aktif bila baris perintah (prompt) kini "
            "diawali tulisan (venv).", bold_lead="Tanda berhasil: ")
para(doc, "Perhatikan bahwa pemisah folder pada Windows memakai garis miring "
          "terbalik (backslash), yaitu venv\\Scripts\\..., berbeda dengan Linux atau "
          "macOS yang memakai garis miring biasa.")

h3(doc, "2.3.3  Memasang Pustaka yang Dibutuhkan")
para(doc, "Setelah virtual environment aktif (prompt diawali (venv)), pustaka "
          "pendukung aplikasi dipasang menggunakan perintah pip, yaitu pengelola "
          "paket bawaan Python. Dua pustaka utama yang dibutuhkan adalah Streamlit "
          "(untuk antarmuka web) dan Pandas (untuk mengolah tabel). Pemasangan dapat "
          "dilakukan langsung dengan menyebut nama pustakanya, atau melalui berkas "
          "daftar pustaka requirements.txt bila tersedia. Pastikan komputer "
          "terhubung ke internet karena pip akan mengunduh pustaka tersebut.")
figure(doc, CODE / "Command_Prompt_memasang_pustaka_Windows.png",
       "Memasang pustaka Streamlit dan Pandas melalui pip di Windows.")

h3(doc, "2.3.4  Menjalankan Aplikasi")
para(doc, "Bila seluruh pustaka sudah terpasang, aplikasi dijalankan dengan perintah "
          "streamlit run app.py dari dalam folder proyek. Streamlit kemudian "
          "menyalakan server lokal dan secara otomatis membuka aplikasi pada "
          "peramban (browser) di alamat http://localhost:8501. Pengguna dapat masuk "
          "menggunakan akun demo admin dengan kata sandi admin123. Selama aplikasi "
          "berjalan, jendela Command Prompt harus tetap dibuka; menutupnya akan "
          "menghentikan aplikasi.")
figure(doc, CODE / "Command_Prompt_menjalankan_aplikasi_Windows.png",
       "Menjalankan aplikasi dengan Streamlit di Windows.")

h3(doc, "2.3.5  Alternatif untuk Linux/macOS")
para(doc, "Bagi pengguna sistem operasi Linux atau macOS, langkahnya serupa namun "
          "perintahnya sedikit berbeda: pembuatan virtual environment memakai "
          "python3 -m venv venv dan pengaktifannya memakai source venv/bin/activate "
          "(dengan garis miring biasa). Setelah aktif, pemasangan pustaka dan "
          "menjalankan aplikasi menggunakan perintah yang sama, yaitu "
          "pip install streamlit pandas dan streamlit run app.py.")
figure(doc, CODE / "Terminal_persiapan_Linux_atau_macOS.png",
       "Persiapan lingkungan pada terminal Linux atau macOS (alternatif).")

# ---- 2.4 ----
h2(doc, "2.4  Struktur Basis Data dan Bagan Akun")
para(doc, "Aplikasi menggunakan basis data SQLite dengan tiga tabel utama yang saling "
          "berelasi. Struktur ini dirancang untuk mendukung pencatatan jurnal "
          "berpasangan, di mana setiap jurnal dapat memiliki beberapa baris akun "
          "(debit dan kredit).")

h3(doc, "2.4.1  Skema Tabel")
para(doc, "Skema tabel dibuat melalui fungsi create_tables(). Terdapat tiga tabel: "
          "users untuk menyimpan data pengguna, jurnal sebagai kepala (header) "
          "transaksi, dan jurnal_baris sebagai rincian akun pada setiap transaksi.")
figure(doc, CODE / "database.py_skema_tabel.png",
       "Definisi skema tabel pada database.py.")
bullet(doc, "menyimpan data akun pengguna meliputi username, kata sandi terenkripsi, dan nama.",
       bold_lead="users: ")
bullet(doc, "menyimpan kepala transaksi berisi kode, tanggal, keterangan, dan tipe jurnal.",
       bold_lead="jurnal: ")
bullet(doc, "menyimpan rincian baris akun (kode akun, nama akun, debit, kredit) yang "
            "merujuk ke tabel jurnal melalui jurnal_id.", bold_lead="jurnal_baris: ")
para(doc, "Nilai uang (debit dan kredit) disimpan dalam tipe TEXT agar presisi angka "
          "tetap terjaga ketika diolah kembali menggunakan tipe Decimal di sisi "
          "aplikasi.")

h3(doc, "2.4.2  Pengisian Data Awal (Seed)")
para(doc, "Saat aplikasi pertama kali dijalankan, basis data diisi dengan data awal "
          "berupa 20 transaksi contoh (T01 sampai T20) dan satu akun pengguna admin. "
          "Hal ini memudahkan pengguna baru untuk langsung melihat contoh laporan yang "
          "telah terisi.")
figure(doc, CODE / "database.py_seed_database.png",
       "Fungsi pengisian data awal basis data.")

h3(doc, "2.4.3  Bagan Akun (Chart of Accounts)")
para(doc, "Aplikasi menggunakan bagan akun yang terdiri atas 15 akun, mencakup "
          "kelompok aset, ekuitas, pendapatan, dan beban. Bagan akun ini menjadi acuan "
          "dalam pencatatan setiap transaksi.")
figure(doc, CODE / "seed_data.py_bagan_akun.png",
       "Definisi bagan akun pada seed_data.py.")
kv_table(doc, [
    ["Aset", "Kas, Perlengkapan, Peralatan, Akumulasi Penyusutan Peralatan"],
    ["Ekuitas", "Modal Pemilik, Prive"],
    ["Pendapatan", "Pendapatan Penjualan Gabah"],
    ["Beban", "Beban Benih, Beban Pupuk, Beban Tenaga Kerja, Beban Pestisida & Obat, "
              "Beban Angkut, Beban Konsumsi, Beban Administrasi, Beban Perlengkapan"],
], headers=["Kelompok Akun", "Contoh Akun"])

# ---- 2.5 ----
h2(doc, "2.5  Implementasi Kode Program")
para(doc, "Bagian ini menjelaskan implementasi kode program inti aplikasi. Setiap "
          "potongan kode disertai penjelasan mengenai fungsi dan logikanya agar mudah "
          "dipahami.")

code_sections = [
    ("2.5.1  Konfigurasi dan Koneksi Aplikasi (app.py)",
     "Berkas app.py merupakan titik awal (entry point) aplikasi, yaitu berkas "
     "pertama yang dijalankan ketika perintah streamlit run app.py dipanggil. Pada "
     "bagian ini dilakukan konfigurasi halaman melalui st.set_page_config untuk "
     "menentukan judul tab, ikon padi, dan tata letak lebar (wide), lalu gaya "
     "tampilan (CSS) bernuansa hijau disuntikkan satu kali agar seluruh aplikasi "
     "tampak seragam dengan tema pertanian. Konfigurasi halaman ini wajib menjadi "
     "perintah Streamlit yang pertama dipanggil; bila tidak, Streamlit akan "
     "menolaknya. Selanjutnya koneksi ke basis data dibuka melalui fungsi get_conn "
     "yang ditandai @st.cache_resource (artinya koneksi dibuat sekali saja dan "
     "dipakai ulang selama aplikasi hidup, bukan dibuka berulang setiap interaksi). "
     "Di dalam get_conn, skema tabel disiapkan, data awal diisi, dan akun admin "
     "dibuat, sehingga aplikasi siap pakai sejak pertama kali dijalankan.",
     "app.py_konfigurasi_dan_koneksi.png",
     "Konfigurasi halaman dan koneksi awal pada app.py.",
     ["st.set_page_config menentukan judul \u201cAkuntansi Tani Padi\u201d, ikon padi, "
      "dan tata letak wide (lebar penuh).",
      "Pemanggilan set_page_config harus berada paling awal sebelum perintah "
      "Streamlit lain agar tidak menimbulkan galat.",
      "ui_helpers.inject_css() menyuntikkan gaya CSS hijau satu kali untuk kartu "
      "statistik, badge, dan tabel.",
      "Fungsi get_conn diberi penanda @st.cache_resource (cache sumber daya) agar "
      "koneksi basis data hanya dibuat sekali dan dipakai ulang.",
      "Di dalam get_conn dipanggil create_tables, seed_database, dan "
      "seed_default_user sehingga tabel, data contoh, dan akun admin langsung "
      "tersedia.",
      "Koneksi memakai check_same_thread=False supaya aman digunakan lintas thread "
      "(proses paralel) oleh Streamlit.",
      "Modul ini sengaja tidak memuat logika akuntansi; tugasnya murni mengatur "
      "konfigurasi, koneksi, dan pengalihan halaman."]),

    ("2.5.2  Menu dan Pengalihan Halaman (app.py)",
     "Setelah pengguna berhasil masuk, aplikasi menampilkan menu navigasi pada "
     "sidebar (panel samping) menggunakan st.radio, yaitu daftar pilihan berbentuk "
     "tombol pilih tunggal. Daftar menu disimpan dalam variabel MENU yang berisi "
     "urutan tetap mulai dari Dashboard, sebelas tahap laporan akuntansi, hingga "
     "menu Input Transaksi. Fungsi _render_sidebar bertugas menampilkan sapaan nama "
     "pengguna, tombol keluar (logout), serta menu tersebut, lalu mengembalikan "
     "pilihan yang ditekan pengguna. Pilihan itu kemudian diteruskan ke fungsi "
     "_dispatch yang mencocokkan setiap pilihan menu dengan fungsi penampil halaman "
     "yang sesuai. Dengan pola terpusat seperti ini, penambahan atau perubahan menu "
     "cukup dilakukan di satu tempat sehingga navigasi tetap rapi dan mudah "
     "dirawat.",
     "app.py_menu_dan_dispatch.png",
     "Logika menu dan pengalihan halaman pada app.py.",
     ["Daftar MENU memuat urutan tetap: Dashboard, 11 laporan siklus akuntansi, dan "
      "Input Transaksi.",
      "st.radio menampilkan menu sebagai pilihan tunggal di sidebar (panel samping "
      "kiri).",
      "Nama pengguna pada sapaan dilewatkan html.escape() agar aman dari "
      "penyisipan kode berbahaya (XSS).",
      "Tombol \u201cKeluar\u201d memanggil st.session_state.clear() lalu st.rerun() "
      "untuk mengakhiri sesi login.",
      "Fungsi _dispatch memetakan tiap teks pilihan menu ke satu fungsi render "
      "halaman tertentu.",
      "Halaman laporan menerima data hasil olahan, sedangkan Dashboard dan Input "
      "menerima koneksi (conn) langsung.",
      "Pemusatan logika menu di satu fungsi memudahkan penambahan menu baru tanpa "
      "mengubah banyak berkas."]),

    ("2.5.3  Alur Utama Aplikasi (app.py)",
     "Fungsi main() adalah pengatur keseluruhan alur aplikasi dan dipanggil di "
     "bagian paling bawah berkas app.py. Pertama, ia memanggil get_conn untuk "
     "memperoleh koneksi basis data yang sudah disiapkan. Kemudian ia memeriksa "
     "status login pengguna yang disimpan pada session_state (memori sesi yang "
     "bertahan selama pengguna membuka aplikasi). Bila pengguna belum masuk, "
     "aplikasi hanya menampilkan halaman autentikasi lalu berhenti dengan st.stop() "
     "sehingga menu dan laporan tidak bisa diakses tanpa login. Sebaliknya, bila "
     "sudah masuk, aplikasi menampilkan sidebar, menyiapkan data laporan melalui "
     "ui_helpers.get_data, lalu menampilkan halaman sesuai pilihan menu. Pola ini "
     "menjadi \u201cgerbang\u201d keamanan yang memastikan hanya pengguna terverifikasi "
     "yang dapat melihat data keuangan.",
     "app.py_alur_utama_main.png",
     "Fungsi utama main() yang mengatur alur aplikasi.",
     ["main() pertama-tama memanggil get_conn() untuk memperoleh koneksi basis data "
      "yang siap pakai.",
      "Status login disimpan pada st.session_state.logged_in dan diinisialisasi "
      "False bila belum ada.",
      "Bila belum login, hanya pages_auth.render_auth yang ditampilkan, diakhiri "
      "st.stop() agar eksekusi berhenti.",
      "st.stop() berfungsi sebagai gerbang keamanan: menu dan laporan tidak dimuat "
      "sebelum pengguna masuk.",
      "Setelah login berhasil, _render_sidebar dipanggil untuk memperoleh pilihan "
      "menu pengguna.",
      "ui_helpers.get_data(conn) menyiapkan seluruh data laporan sebelum halaman "
      "dirender.",
      "Pemisahan tegas antara keadaan \u201cbelum login\u201d dan \u201csudah login\u201d "
      "membuat alur mudah dipahami dan diaudit."]),

    ("2.5.4  Pengamanan Kata Sandi (auth.py)",
     "Berkas auth.py menangani keamanan akun pengguna. Kata sandi tidak pernah "
     "disimpan dalam bentuk aslinya, melainkan diubah menjadi hash (sidik jari "
     "digital satu arah yang tidak bisa dikembalikan ke teks asli) menggunakan "
     "algoritme PBKDF2-HMAC-SHA256. Fungsi hash_password membangkitkan salt acak "
     "(potongan data acak unik per pengguna) sepanjang 16 byte, lalu mengolah kata "
     "sandi sebanyak 200.000 iterasi (pengulangan) agar sulit ditebak walau dengan "
     "komputer cepat. Hasilnya disimpan dalam satu teks berformat "
     "pbkdf2_sha256$<iterasi>$<salt>$<hash> sehingga seluruh informasi verifikasi "
     "tersimpan rapi tanpa membocorkan kata sandi. Saat login, verifikasi_password "
     "memproses kata sandi yang dimasukkan dengan salt dan iterasi yang sama, lalu "
     "membandingkan hasilnya memakai secrets.compare_digest yang aman terhadap "
     "serangan pengukuran waktu (timing attack).",
     "auth.py_hash_dan_verifikasi_password.png",
     "Fungsi hash dan verifikasi kata sandi pada auth.py.",
     ["Konstanta ITERASI bernilai 200.000, menentukan berapa kali kata sandi diolah "
      "agar hash sulit dibongkar.",
      "secrets.token_bytes(16) menghasilkan salt acak unik untuk tiap pengguna "
      "sehingga dua sandi sama tetap berbeda hash-nya.",
      "hashlib.pbkdf2_hmac(\"sha256\", ...) adalah inti perhitungan hash satu arah "
      "berbasis SHA-256.",
      "Hasil disimpan sebagai teks pbkdf2_sha256$<iterasi>$<salt>$<hash> agar mudah "
      "dipisah kembali saat verifikasi.",
      "verifikasi_password memecah teks tersimpan, menghitung ulang hash dari kata "
      "sandi masukan, lalu membandingkannya.",
      "Perbandingan memakai secrets.compare_digest yang aman dari timing attack "
      "(penebakan lewat selisih waktu proses).",
      "Pendekatan ini memakai pustaka bawaan Python (hashlib + secrets) tanpa "
      "dependensi tambahan, sehingga ringan dan andal."]),

    ("2.5.5  Pendaftaran Pengguna (auth.py)",
     "Fungsi register bertugas menambahkan pengguna baru ke basis data dengan "
     "beberapa lapis pemeriksaan agar data tetap bersih. Mula-mula username "
     "dirapikan dari spasi berlebih, lalu dipastikan username dan kata sandi tidak "
     "kosong. Kata sandi juga diwajibkan minimal enam karakter sebagai syarat "
     "keamanan dasar. Sebelum menyimpan, fungsi memeriksa apakah username sudah "
     "dipakai; bila sudah ada, pendaftaran ditolak agar tidak terjadi duplikasi. "
     "Bila semua syarat terpenuhi, kata sandi dienkripsi memakai hash_password, "
     "kemudian data dimasukkan ke tabel users dan disimpan permanen. Fungsi "
     "mengembalikan pasangan nilai (berhasil/gagal beserta pesannya) sehingga "
     "halaman pendaftaran dapat menampilkan umpan balik yang jelas kepada pengguna.",
     "auth.py_register.png",
     "Fungsi pendaftaran pengguna baru pada auth.py.",
     ["Username dirapikan dengan .strip() untuk menghapus spasi di awal dan akhir.",
      "Username dan kata sandi wajib diisi; bila kosong, pendaftaran langsung "
      "ditolak dengan pesan jelas.",
      "Kata sandi diwajibkan minimal 6 karakter sebagai pengaman dasar.",
      "Perintah SELECT 1 FROM users WHERE username = ? memeriksa apakah username "
      "sudah terpakai sebelum menyimpan.",
      "Tanda tanya (?) pada perintah SQL adalah parameter aman yang mencegah "
      "serangan SQL injection (penyisipan perintah jahat).",
      "Kata sandi disimpan sebagai hash hasil hash_password, bukan teks asli.",
      "Fungsi mengembalikan tuple (bool, pesan) agar antarmuka tahu pendaftaran "
      "berhasil atau gagal beserta alasannya."]),

    ("2.5.6  Formulir Login (pages_auth.py)",
     "Berkas pages_auth.py menyajikan halaman masuk dan daftar dalam dua tab "
     "(\u201cMasuk\u201d dan \u201cDaftar\u201d) di kolom tengah layar. Fungsi render_login "
     "menampilkan formulir berisi kolom username dan kata sandi, dengan kolom kata "
     "sandi disembunyikan (type=\"password\"). Modul ini bersifat read-only "
     "terhadap backend, artinya ia tidak menghitung hash sendiri melainkan hanya "
     "memanggil auth.login untuk memverifikasi kredensial. Apabila login berhasil, "
     "informasi pengguna (id, username, nama) disimpan ke session_state dan tanda "
     "logged_in dijadikan True, lalu halaman dimuat ulang dengan st.rerun. Bila "
     "gagal, pesan kesalahan ditampilkan tanpa membocorkan apakah username atau "
     "kata sandinya yang salah, sehingga lebih aman.",
     "pages_auth.py_form_login.png",
     "Formulir login pada pages_auth.py.",
     ["st.form(\"form_login\") membungkus kolom input agar diproses sekaligus saat "
      "tombol ditekan.",
      "Kolom kata sandi memakai type=\"password\" sehingga karakter tersembunyi "
      "saat diketik.",
      "Verifikasi kredensial didelegasikan ke auth.login; halaman tidak pernah "
      "menghitung hash sendiri.",
      "Jika berhasil, info pengguna disimpan ke st.session_state.user dan "
      "logged_in di-set True.",
      "st.rerun() memuat ulang halaman agar tampilan langsung berpindah ke menu "
      "utama.",
      "Pesan kesalahan dibuat umum (\u201cUsername atau password salah\u201d) agar tidak "
      "membocorkan akun mana yang ada.",
      "Hanya data non-sensitif (id/username/nama) yang disimpan; kata sandi polos "
      "tidak pernah ikut tersimpan."]),

    ("2.5.7  Penomoran Kode dan Penyimpanan Jurnal (database.py)",
     "Bagian ini berisi dua fungsi yang bekerja sama saat sebuah transaksi baru "
     "disimpan. Fungsi next_kode menghitung jumlah jurnal umum yang sudah ada lalu "
     "menambah satu, menghasilkan kode urut berformat T## (misalnya T21 setelah "
     "T20). Fungsi insert_jurnal menyimpan transaksi tersebut, tetapi sebelum "
     "menulis apa pun ia memanggil acc.validasi_entry untuk memastikan total debit "
     "sama dengan total kredit; bila tidak seimbang, fungsi melempar ValueError "
     "sehingga tidak ada data yang tersimpan. Setelah lolos validasi, kepala jurnal "
     "(tanggal, keterangan, tipe) dimasukkan ke tabel jurnal, lalu setiap baris "
     "akun disimpan ke tabel jurnal_baris yang terhubung melalui jurnal_id. Nilai "
     "uang diubah ke Decimal lalu disimpan sebagai teks agar presisinya terjaga, "
     "dan seluruh proses ditutup dengan conn.commit() agar perubahan permanen.",
     "database.py_next_kode_dan_insert_jurnal.png",
     "Fungsi penomoran kode dan penyimpanan jurnal pada database.py.",
     ["next_kode menghitung COUNT jurnal umum lalu menambah 1, menghasilkan kode "
      "T## (contoh T21) secara otomatis.",
      "Format f\"T{n:02d}\" memastikan nomor selalu dua digit (T01, T02, ... T21).",
      "insert_jurnal memanggil acc.validasi_entry lebih dulu; bila debit tidak sama "
      "dengan kredit, ValueError dilempar.",
      "Karena validasi dilakukan sebelum menulis, transaksi tidak seimbang tidak "
      "akan pernah tersimpan ke basis data.",
      "Kepala transaksi disimpan ke tabel jurnal, lalu cur.lastrowid dipakai "
      "sebagai jurnal_id penghubung ke baris akun.",
      "Tiap baris akun disimpan ke jurnal_baris dengan nilai uang diubah ke Decimal "
      "lalu disimpan sebagai teks demi presisi.",
      "conn.commit() di akhir memastikan seluruh penyimpanan benar-benar permanen "
      "(tersimpan tetap)."]),

    ("2.5.8  Membaca Jurnal Umum (database.py)",
     "Fungsi pembacaan jurnal bertugas mengambil data transaksi dari basis data "
     "untuk ditampilkan maupun diolah menjadi laporan. Fungsi inti _baca_jurnal "
     "membaca kepala jurnal dari tabel jurnal, lalu untuk setiap kepala tersebut "
     "membaca baris-baris akunnya dari tabel jurnal_baris. Data disusun menjadi "
     "struktur bertingkat (dictionary): satu transaksi memiliki daftar baris akun "
     "di dalamnya. Saat dibaca, nilai uang yang tersimpan sebagai teks diubah "
     "kembali menjadi Decimal agar perhitungan tetap akurat. Fungsi get_jurnal_umum "
     "memanggil _baca_jurnal khusus untuk tipe \u201cumum\u201d sehingga hanya transaksi "
     "operasional (termasuk yang ditambah pengguna) yang diambil, sedangkan "
     "get_semua_jurnal mengambil seluruh tipe jurnal.",
     "database.py_get_jurnal_umum_dan_baca_jurnal.png",
     "Fungsi pembacaan jurnal umum pada database.py.",
     ["_baca_jurnal membaca kepala jurnal dahulu, lalu baris akun tiap transaksi "
      "secara berurutan (ORDER BY id).",
      "Parameter tipe (umum/penyesuaian/penutup) menyaring jurnal yang dibaca "
      "sesuai kebutuhan halaman.",
      "Nilai debit dan kredit yang tersimpan TEXT diubah menjadi Decimal saat "
      "dibaca agar tidak ada galat pembulatan.",
      "Setiap transaksi dikembalikan sebagai dictionary berisi db_id, kode, "
      "tanggal, keterangan, tipe, dan daftar lines.",
      "get_jurnal_umum memanggil _baca_jurnal dengan tipe=\"umum\" untuk transaksi "
      "operasional sehari-hari.",
      "get_semua_jurnal memanggil tanpa filter tipe sehingga mengambil semua jenis "
      "jurnal sekaligus.",
      "Struktur bertingkat ini memudahkan accounting.py mengolah data menjadi buku "
      "besar dan laporan turunannya."]),

    ("2.5.9  Mengubah Jurnal (database.py)",
     "Fungsi update_jurnal memungkinkan pengguna memperbaiki transaksi yang sudah "
     "tersimpan tanpa harus menghapus dan membuatnya ulang. Sama seperti saat "
     "menyimpan transaksi baru, fungsi ini terlebih dahulu memanggil "
     "acc.validasi_entry untuk memastikan debit dan kredit tetap seimbang; bila "
     "tidak, ValueError dilempar dan data lama tetap utuh. Setelah lolos validasi, "
     "fungsi memperbarui kepala jurnal (tanggal dan keterangan) pada tabel jurnal. "
     "Untuk rinciannya, pendekatan yang dipakai adalah menghapus seluruh baris akun "
     "lama lalu memasukkan baris akun yang baru, sehingga rincian selalu konsisten "
     "dengan masukan terbaru. Kode dan tipe jurnal sengaja tidak diubah agar "
     "identitas transaksi tetap terjaga, dan seluruh perubahan dikunci permanen "
     "dengan conn.commit().",
     "database.py_update_jurnal.png",
     "Fungsi pembaruan jurnal pada database.py.",
     ["acc.validasi_entry dipanggil di awal sehingga pembaruan tidak seimbang "
      "ditolak dan data lama tetap aman.",
      "Perintah UPDATE jurnal SET tanggal, keterangan memperbarui kepala transaksi "
      "sesuai masukan baru.",
      "Seluruh baris akun lama dihapus dengan DELETE FROM jurnal_baris WHERE "
      "jurnal_id = ?.",
      "Baris akun baru kemudian dimasukkan ulang, menjamin rincian selalu cocok "
      "dengan data terbaru.",
      "Kode dan tipe jurnal tidak diubah agar identitas transaksi (misalnya T21) "
      "tetap konsisten.",
      "Nilai uang tetap diubah ke Decimal lalu disimpan sebagai teks, sama seperti "
      "saat penyimpanan awal.",
      "conn.commit() memastikan pembaruan kepala dan rincian tersimpan sebagai satu "
      "kesatuan yang utuh."]),

    ("2.5.10  Menghapus dan Mereset Jurnal (database.py)",
     "Bagian ini memuat dua fungsi pengelolaan data, yaitu hapus_jurnal untuk "
     "menghapus satu transaksi dan reset_jurnal untuk mengembalikan seluruh data ke "
     "kondisi awal. Fungsi hapus_jurnal mula-mula menghapus seluruh baris akun "
     "milik transaksi tersebut dari tabel jurnal_baris, baru kemudian menghapus "
     "kepala transaksinya dari tabel jurnal, sehingga tidak ada baris akun yatim "
     "(baris yang induknya sudah tidak ada). Fungsi reset_jurnal mengosongkan kedua "
     "tabel jurnal lalu memanggil seed_database untuk mengisi kembali 20 transaksi "
     "contoh. Reset ini berguna ketika pengguna ingin berlatih dari awal atau "
     "membersihkan transaksi percobaan. Karena bersifat menghapus banyak data, "
     "fitur reset di antarmuka diberi konfirmasi terlebih dahulu agar tidak terpicu "
     "tanpa sengaja.",
     "database.py_hapus_jurnal_dan_reset_jurnal.png",
     "Fungsi penghapusan dan reset jurnal pada database.py.",
     ["hapus_jurnal menghapus baris akun lebih dulu (jurnal_baris) baru kepala "
      "transaksinya (jurnal).",
      "Urutan ini mencegah munculnya baris akun yatim, yaitu rincian yang induk "
      "transaksinya sudah terhapus.",
      "reset_jurnal menjalankan DELETE pada jurnal_baris dan jurnal untuk "
      "mengosongkan seluruh data transaksi.",
      "Setelah dikosongkan, reset_jurnal memanggil seed_database untuk mengisi "
      "ulang 20 transaksi contoh (T01-T20).",
      "Reset bermanfaat untuk berlatih dari nol atau menghapus transaksi percobaan "
      "yang menumpuk.",
      "Karena menghapus banyak data sekaligus, fitur reset di antarmuka memakai "
      "kotak konfirmasi sebelum dijalankan.",
      "conn.commit() pada kedua fungsi memastikan penghapusan benar-benar tersimpan "
      "ke berkas basis data."]),

    ("2.5.11  Menyimpan Transaksi dari Formulir (pages_input.py)",
     "Berkas pages_input.py adalah halaman tempat pengguna mencatat transaksi baru, "
     "dan ia dirancang dengan pengaman ketat. Formulir input meminta tanggal, "
     "keterangan, satu baris akun debit beserta nominalnya, dan satu baris akun "
     "kredit beserta nominalnya. Akun tidak diketik bebas melainkan dipilih dari "
     "bagan akun resmi (chart of accounts) melalui st.selectbox, sehingga pengguna "
     "tidak mungkin memasukkan akun di luar 15 akun yang sah. Ketika tombol simpan "
     "ditekan, kedua baris dirakit menjadi daftar lines lalu dikirim ke "
     "db.insert_jurnal yang memvalidasi keseimbangan debit-kredit. Bila tidak "
     "seimbang, pesan kesalahan ditampilkan dan tidak ada yang tersimpan; bila "
     "berhasil, muncul notifikasi sukses dan halaman dimuat ulang agar daftar "
     "transaksi langsung diperbarui.",
     "pages_input.py_simpan_transaksi.png",
     "Proses penyimpanan transaksi pada pages_input.py.",
     ["Formulir dibungkus st.form dengan clear_on_submit=True agar kolom otomatis "
      "kosong setelah transaksi tersimpan.",
      "Akun debit dan kredit dipilih dari bagan akun lewat st.selectbox, bukan "
      "diketik bebas, sehingga selalu sah.",
      "Nominal dimasukkan lewat st.number_input dengan langkah (step) 1000 dan nilai "
      "minimal 0.",
      "Dua baris (debit dan kredit) dirakit menjadi daftar lines berisi kode, nama "
      "akun, debit, dan kredit.",
      "db.insert_jurnal dipanggil di dalam blok try; ValueError dari validasi "
      "ditangkap dan ditampilkan sebagai st.error.",
      "Bila tidak seimbang, tidak ada data tersimpan; bila berhasil, st.success "
      "muncul dan st.rerun memuat ulang daftar.",
      "Pembatasan akun ke bagan resmi menjaga konsistensi laporan karena setiap "
      "akun sudah punya tipe dan saldo normal yang jelas."]),

    ("2.5.12  Tombol Ubah dan Hapus (pages_input.py)",
     "Pada daftar transaksi, setiap baris dilengkapi tombol untuk mengubah atau "
     "menghapus agar pengelolaan data terasa mudah. Untuk menjaga keutuhan data "
     "contoh, transaksi seed (T01-T20) dikunci dan hanya diberi label "
     "\u201cdata awal\u201d tanpa tombol, sedangkan transaksi tambahan buatan pengguna "
     "memiliki tombol Edit dan Hapus. Ketika tombol Edit ditekan, db_id transaksi "
     "terpilih disimpan ke session_state.edit_db_id sehingga aplikasi tahu "
     "transaksi mana yang sedang disunting setelah halaman dimuat ulang. Ketika "
     "tombol Hapus ditekan, fungsi db.hapus_jurnal langsung dipanggil untuk "
     "menghapus transaksi tersebut. Pemisahan antara data terkunci dan data yang "
     "bisa diubah ini melindungi contoh bawaan agar tidak terhapus tanpa sengaja.",
     "pages_input.py_tombol_edit_dan_hapus.png",
     "Tombol ubah dan hapus transaksi pada pages_input.py.",
     ["Himpunan kode seed (T01-T20) dipakai untuk membedakan transaksi contoh dari "
      "transaksi tambahan pengguna.",
      "Transaksi seed ditandai sebagai \u201cdata awal\u201d (terkunci) dan tidak diberi "
      "tombol agar tidak bisa diubah atau dihapus.",
      "Tiap baris ditata dalam tiga kolom (informasi, tombol edit, tombol hapus) "
      "memakai st.columns.",
      "Tombol Edit menyimpan db_id transaksi ke st.session_state.edit_db_id lalu "
      "memanggil st.rerun.",
      "session_state (memori sesi) menjaga transaksi mana yang sedang disunting "
      "meski halaman dimuat ulang.",
      "Tombol Hapus memanggil db.hapus_jurnal(conn, db_id) untuk menghapus "
      "transaksi beserta seluruh barisnya.",
      "Setiap tombol diberi key unik (mis. edit_<db_id>) agar Streamlit tidak "
      "mencampuradukkan tombol antar-transaksi."]),

    ("2.5.13  Formulir Ubah Transaksi (pages_input.py)",
     "Ketika sebuah transaksi tambahan dipilih untuk diubah, fungsi _render_form_edit "
     "menampilkan formulir yang sudah terisi (pre-fill) dengan nilai transaksi saat "
     "ini. Fungsi ini cerdas mengenali baris mana yang debit dan mana yang kredit, "
     "lalu menempatkan akun yang benar pada masing-masing selectbox sebagai pilihan "
     "awal. Tanggal lama juga dimuat kembali, dengan pengaman bila format tanggalnya "
     "tidak valid maka dipakai tanggal hari ini. Pengguna dapat menyunting tanggal, "
     "keterangan, akun, dan nominal, lalu menekan Simpan Perubahan atau Batal. Saat "
     "disimpan, data dikirim ke db.update_jurnal yang kembali memvalidasi "
     "keseimbangan; bila gagal ditampilkan pesan error, bila berhasil mode edit "
     "ditutup dan halaman dimuat ulang. Tombol Batal sekadar mengosongkan status "
     "edit tanpa menyimpan apa pun.",
     "pages_input.py_form_edit_transaksi.png",
     "Formulir ubah transaksi pada pages_input.py.",
     ["Baris debit dan kredit dikenali otomatis dengan mencari baris yang nilai "
      "debit atau kreditnya lebih dari nol.",
      "Akun yang sedang dipakai ditempatkan sebagai pilihan awal selectbox lewat "
      "perhitungan index yang cocok.",
      "Tanggal lama dimuat dengan datetime.date.fromisoformat; bila formatnya rusak "
      "dipakai tanggal hari ini sebagai cadangan.",
      "Setiap kolom input diberi key yang mengandung db_id agar formulir antar "
      "transaksi tidak tertukar.",
      "Tombol \u201cSimpan Perubahan\u201d merakit ulang baris lalu memanggil "
      "db.update_jurnal di dalam blok try.",
      "Jika validasi gagal, st.error menampilkan pesan; jika berhasil, edit_db_id "
      "dikosongkan dan st.rerun dipanggil.",
      "Tombol \u201cBatal\u201d hanya menyetel edit_db_id ke None sehingga formulir "
      "tertutup tanpa mengubah data."]),

    ("2.5.14  Format Rupiah dan Validasi Entry (accounting.py)",
     "Berkas accounting.py adalah mesin akuntansi inti aplikasi, dan bagian ini "
     "memuat dua fungsi pendukung yang sering dipakai. Fungsi format_rupiah mengubah "
     "angka (Decimal atau int) menjadi teks mata uang gaya Indonesia, misalnya nilai "
     "17500000 menjadi \u201cRp 17.500.000\u201d, lengkap dengan tanda titik pemisah "
     "ribuan dan dukungan nilai negatif. Fungsi validasi_entry adalah penjaga "
     "integritas transaksi: ia menjumlahkan seluruh debit dan seluruh kredit pada "
     "satu transaksi, lalu memastikan keduanya sama dan tidak nol. Bila debit tidak "
     "sama dengan kredit, atau total bernilai nol, fungsi melempar ValueError "
     "dengan pesan yang menjelaskan masalahnya. Karena validasi ini dipanggil "
     "sebelum setiap penyimpanan dan pembaruan, prinsip dasar akuntansi berpasangan "
     "(debit = kredit) selalu terjaga.",
     "accounting.py_format_rupiah_dan_validasi_entry.png",
     "Fungsi format rupiah dan validasi entry pada accounting.py.",
     ["format_rupiah mengubah Decimal/int menjadi teks seperti \u201cRp 17.500.000\u201d "
      "dengan titik sebagai pemisah ribuan.",
      "Tanda titik diperoleh dengan mengganti koma pemformatan (f\"{abs(n):,}\") "
      "menjadi titik agar sesuai gaya Indonesia.",
      "format_rupiah juga menangani nilai negatif dengan menambahkan tanda minus di "
      "depan angka.",
      "validasi_entry menjumlahkan seluruh debit dan seluruh kredit memakai tipe "
      "Decimal agar presisi.",
      "Bila total debit tidak sama dengan total kredit, ValueError dilempar beserta "
      "pesan selisihnya.",
      "Bila total bernilai nol (transaksi kosong), ValueError juga dilempar agar "
      "tidak ada jurnal hampa tersimpan.",
      "Karena dipanggil sebelum setiap simpan dan ubah, fungsi ini menjamin prinsip "
      "akuntansi berpasangan (debit = kredit) selalu terpenuhi."]),

    ("2.5.15  Penyusunan Laporan Laba Rugi (accounting.py)",
     "Fungsi laba_rugi menghasilkan salah satu laporan terpenting bagi petani, "
     "yaitu seberapa besar untung atau rugi usaha tani pada satu periode. Fungsi ini "
     "pertama-tama membangun buku besar dari daftar jurnal agar saldo tiap akun "
     "diketahui. Kemudian ia menelusuri seluruh akun: untuk akun bertipe pendapatan, "
     "nilainya dihitung dari selisih kredit dikurangi debit; untuk akun bertipe "
     "beban, nilainya dihitung dari selisih debit dikurangi kredit. Seluruh "
     "pendapatan dan seluruh beban dijumlahkan masing-masing, lalu laba bersih "
     "diperoleh dari total pendapatan dikurangi total beban. Hasilnya dikembalikan "
     "sebagai dictionary berisi rincian baris pendapatan, baris beban, totalnya, "
     "serta laba bersih, sehingga halaman laporan tinggal menampilkannya dengan "
     "rapi.",
     "accounting.py_laba_rugi.png",
     "Fungsi penyusunan laporan laba rugi pada accounting.py.",
     ["Fungsi memulai dengan membangun buku_besar(jurnal) agar saldo setiap akun "
      "tersedia lebih dulu.",
      "Akun diurutkan berdasarkan kode akun supaya urutan baris laporan konsisten "
      "dan mudah dibaca.",
      "Untuk akun pendapatan, nilai dihitung sebagai total_kredit dikurangi "
      "total_debit (saldo normal kredit).",
      "Untuk akun beban, nilai dihitung sebagai total_debit dikurangi total_kredit "
      "(saldo normal debit).",
      "total_pendapatan dan total_beban dijumlahkan terpisah memakai tipe Decimal "
      "agar tetap presisi.",
      "laba_bersih dihitung sebagai total_pendapatan dikurangi total_beban, "
      "menunjukkan untung (positif) atau rugi (negatif).",
      "Hasil dikemas dalam dictionary (baris pendapatan, baris beban, total, dan "
      "laba bersih) agar mudah ditampilkan halaman laporan."]),

    ("2.5.16  Pengambilan Data Laporan (ui_helpers.py)",
     "Fungsi get_data pada ui_helpers.py berperan sebagai satu sumber kebenaran "
     "(single source of truth) untuk seluruh data laporan, sehingga semua halaman "
     "memakai data yang sama dan konsisten. Pertama, ia membaca jurnal umum dari "
     "basis data melalui db.get_jurnal_umum. Lalu ia menyusun versi "
     "\u201cdisesuaikan\u201d dengan menggabungkan jurnal umum dan jurnal penyesuaian. "
     "Dari versi disesuaikan itu, ia menghitung jurnal penutup memakai "
     "acc.jurnal_penutup, dan terakhir menyusun versi \u201csetelah penutupan\u201d "
     "dengan menggabungkan jurnal disesuaikan dan jurnal penutup. Keempat himpunan "
     "data ini dikembalikan dalam satu dictionary, sehingga setiap halaman laporan "
     "tinggal mengambil bagian yang dibutuhkannya tanpa menghitung ulang dari awal.",
     "ui_helpers.py_get_data.png",
     "Fungsi get_data untuk menyiapkan data laporan pada ui_helpers.py.",
     ["get_data membaca jurnal umum dari basis data sebagai titik awal seluruh "
      "perhitungan laporan.",
      "Versi \u201cdisesuaikan\u201d dibentuk dengan menambahkan jurnal penyesuaian ke "
      "jurnal umum.",
      "Jurnal penutup dihitung dari versi disesuaikan memakai acc.jurnal_penutup.",
      "Versi \u201csetelah penutupan\u201d dibentuk dari jurnal disesuaikan ditambah "
      "jurnal penutup.",
      "Keempat himpunan (jurnal, disesuaikan, penutup, setelah_penutupan) "
      "dikembalikan dalam satu dictionary.",
      "Modul ini bersifat read-only: ia hanya membaca dan merangkai data, tidak "
      "pernah menulis ke basis data.",
      "Dengan satu sumber data terpusat, seluruh halaman laporan dijamin "
      "menampilkan angka yang konsisten satu sama lain."]),

    ("2.5.17  Penyajian Tabel Berstyle (ui_helpers.py)",
     "Agar setiap laporan tampil rapi dan seragam, ui_helpers.py menyediakan fungsi "
     "tabel_html yang mengubah data menjadi tabel HTML berstyle. Fungsi ini menerima "
     "data baik dalam bentuk DataFrame pandas maupun list-of-dict, lalu membangun "
     "bagian kepala (thead) dan badan (tbody) tabel. Kolom angka dapat ditandai "
     "sebagai right_cols sehingga dirata-kanan dan memakai angka tabular agar "
     "kolom rupiah mudah dibandingkan secara vertikal. Demi keamanan, setiap isi sel "
     "dilewatkan html.escape() untuk mencegah serangan XSS (penyisipan kode "
     "berbahaya), terutama pada kolom Keterangan yang berisi teks bebas dari "
     "pengguna. Fungsi juga dapat menandai baris total agar dicetak tebal, sehingga "
     "ringkasan laporan langsung terlihat menonjol dengan warna tema hijau yang "
     "konsisten.",
     "ui_helpers.py_tabel_html.png",
     "Fungsi penyajian tabel HTML berstyle pada ui_helpers.py.",
     ["tabel_html menerima data berupa DataFrame pandas maupun list-of-dict, lalu "
      "menyusunnya jadi tabel HTML.",
      "Parameter right_cols menandai kolom angka agar dirata-kanan dengan class "
      ".ajp-num (angka tabular, tidak terpotong).",
      "Perataan kanan membuat kolom rupiah mudah dibandingkan dari atas ke bawah.",
      "Setiap sel header dan isi dilewatkan html.escape() untuk mencegah XSS "
      "(penyisipan kode berbahaya).",
      "Kolom Keterangan memuat teks bebas dari pengguna, sehingga escape ini "
      "menjadi pengaman penting.",
      "Parameter total_row_label menandai baris total dengan class .ajp-total "
      "sehingga dicetak tebal dan menonjol.",
      "Tabel dibungkus div .ajp-table-wrap sehingga mewarisi gaya tema hijau yang "
      "seragam di seluruh laporan."]),
]

for title, intro, img, cap, bullets in code_sections:
    h3(doc, title)
    para(doc, intro)
    figure(doc, CODE / img, cap)
    para(doc, "Penjelasan:", bold=True, justify=False)
    for b in bullets:
        bullet(doc, b)
    para(doc, "Singkatnya, bagian kode ini menjadi salah satu fondasi penting agar "
              "aplikasi Akuntansi Tani Padi bekerja dengan benar, aman, dan mudah "
              "digunakan oleh pelaku usaha tani.")

# ---- 2.6 ----
h2(doc, "2.6  Panduan Penggunaan Aplikasi")
para(doc, "Bagian ini menjelaskan langkah penggunaan aplikasi disertai tangkapan layar "
          "dari aplikasi yang sedang berjalan. Pengguna dapat mengikuti urutan berikut "
          "untuk memanfaatkan seluruh fitur aplikasi.")

shot_sections = [
    ("2.6.1  Halaman Login",
     "Halaman pertama yang muncul saat aplikasi dibuka adalah halaman Login. Di tengah "
     "layar terdapat form berjudul \u201cMasuk\u201d yang berisi dua kolom isian, yaitu "
     "Username dan Password (kata sandi yang diketik otomatis disamarkan menjadi titik-"
     "titik agar tidak terbaca orang lain). Tepat di bawahnya ada tombol \u201cMasuk\u201d, "
     "serta sebuah tab \u201cDaftar\u201d bagi pengguna yang belum memiliki akun. Ketika "
     "tombol Masuk ditekan, aplikasi mencocokkan data yang diketik dengan data akun yang "
     "tersimpan di basis data; bila cocok, pengguna langsung dibawa masuk ke aplikasi dan "
     "namanya disimpan untuk ditampilkan di dalam aplikasi, sedangkan bila tidak cocok "
     "akan muncul pesan \u201cUsername atau password salah.\u201d Yang penting dipahami, "
     "kata sandi tidak pernah disimpan dalam bentuk teks biasa: sandi yang diketik diubah "
     "dahulu menjadi kode acak (di-hash dengan metode PBKDF2-HMAC-SHA256, 200.000 putaran "
     "dan garam acak unik per pengguna), lalu kode itulah yang dibandingkan dengan kode "
     "yang tersimpan, sehingga sandi asli tetap aman dan tidak pernah dibaca secara polos.",
     "01-login.png", "Halaman login aplikasi.",
     ["Form \u201cMasuk\u201d berisi dua kolom: Username dan Password (sandi tampil sebagai "
      "titik-titik agar terlindung dari mata orang lain).",
      "Tombol \u201cMasuk\u201d memeriksa kecocokan akun ke basis data; bila benar pengguna "
      "langsung masuk, bila salah muncul pesan \u201cUsername atau password salah.\u201d",
      "Tersedia tab \u201cDaftar\u201d untuk berpindah ke halaman pendaftaran bagi yang "
      "belum punya akun.",
      "Untuk mencoba aplikasi, sudah disediakan akun demo dengan username admin dan sandi "
      "admin123.",
      "Sandi yang diketik tidak dibandingkan apa adanya, melainkan diubah menjadi kode "
      "acak terlebih dahulu lalu dicocokkan dengan kode tersimpan (verifikasi aman).",
      "Sesi login disimpan di memori sementara; bila halaman di-refresh, sesi dianggap "
      "baru dan pengguna kembali ke halaman login."]),
    ("2.6.2  Halaman Pendaftaran",
     "Pengguna yang belum memiliki akun dapat membuatnya melalui tab \u201cDaftar\u201d. "
     "Pada halaman ini ditampilkan sebuah form pendaftaran yang berisi kolom Username, "
     "kolom Password yang dilengkapi tulisan bantuan \u201cMinimal 6 karakter.\u201d, "
     "kolom Nama lengkap, dan sebuah tombol \u201cDaftar\u201d di bagian bawah. Setelah "
     "semua kolom diisi dan tombol Daftar ditekan, aplikasi akan memeriksa apakah data "
     "sudah memenuhi syarat. Bila pendaftaran berhasil, muncul pesan sukses yang meminta "
     "pengguna kembali login menggunakan akun barunya; sebaliknya bila gagal, misalnya "
     "username sudah dipakai orang lain atau sandi kurang dari enam karakter, aplikasi "
     "menampilkan pesan kesalahan yang menjelaskan masalahnya. Di balik layar, akun baru "
     "disimpan ke basis data lengkap dengan sandi yang sudah diamankan menjadi kode acak "
     "beserta garam acak unik, sehingga walaupun seseorang membuka isi basis data, sandi "
     "asli tetap tidak bisa terbaca.",
     "02-register.png", "Halaman pendaftaran pengguna baru.",
     ["Form pendaftaran berisi tiga kolom isian: Username, Password (dengan petunjuk "
      "\u201cMinimal 6 karakter.\u201d), dan Nama lengkap.",
      "Tombol \u201cDaftar\u201d menyimpan akun baru setelah memastikan data memenuhi "
      "syarat.",
      "Bila berhasil, muncul pesan sukses yang mengarahkan pengguna untuk login dengan "
      "akun barunya.",
      "Bila gagal, misalnya username sudah terpakai atau sandi kurang dari enam karakter, "
      "muncul pesan kesalahan yang menjelaskan penyebabnya.",
      "Sandi disimpan dalam bentuk kode acak (hash) berbumbu garam unik per pengguna, "
      "bukan teks biasa, sehingga isi basis data tetap aman.",
      "Setelah akun jadi, pengguna kembali ke tab \u201cMasuk\u201d untuk login seperti "
      "biasa."]),
    ("2.6.3  Dasbor Ringkasan",
     "Begitu berhasil login, pengguna disambut halaman Dashboard yang diawali banner "
     "\u201c\U0001F33E Dashboard\u201d dan sapaan \u201cSelamat datang kembali, {nama}.\u201d "
     "sesuai nama pengguna. Pada baris pertama tampil empat kartu metrik, yaitu "
     "\u201cTotal Pendapatan\u201d (\U0001F33E), \u201cTotal Beban\u201d (\U0001F4B8), "
     "\u201cLaba Bersih\u201d (\U0001F4C8), dan \u201cModal Akhir\u201d (\U0001F4B0). Pada "
     "baris kedua tampil dua kartu tambahan, yaitu \u201cKas Akhir\u201d (\U0001F3E6) dan "
     "\u201cJumlah Transaksi\u201d (\U0001F4CB) yang menyebutkan banyaknya transaksi, "
     "misalnya \u201c21 transaksi\u201d. Seluruh angka pada kartu-kartu ini bukan diketik "
     "manual, melainkan dihitung otomatis dari jurnal yang sudah disesuaikan: nilai "
     "diambil dari hasil perhitungan laba rugi, laporan perubahan ekuitas, dan laporan "
     "arus kas, lalu dirangkum ke dalam kartu agar pengguna langsung melihat kondisi "
     "keuangan terkini dalam sekejap tanpa perlu membuka satu per satu laporannya.",
     "03-dashboard.png", "Dasbor ringkasan keuangan.",
     ["Baris pertama memuat empat kartu metrik: \u201cTotal Pendapatan\u201d, \u201cTotal "
      "Beban\u201d, \u201cLaba Bersih\u201d, dan \u201cModal Akhir\u201d.",
      "Baris kedua memuat dua kartu: \u201cKas Akhir\u201d dan \u201cJumlah Transaksi\u201d "
      "(yang menyebutkan jumlah transaksi tercatat).",
      "Banner \u201c\U0001F33E Dashboard\u201d dan sapaan personal menampilkan nama "
      "pengguna yang sedang login.",
      "Semua angka dihitung otomatis dari jurnal yang telah disesuaikan, bukan diketik "
      "secara manual.",
      "Nilai kartu bersumber dari perhitungan laba rugi, perubahan ekuitas, dan arus kas "
      "sehingga selalu konsisten dengan laporan rincinya.",
      "Dasbor berfungsi sebagai ringkasan cepat agar kondisi keuangan langsung terbaca "
      "tanpa membuka setiap laporan satu per satu."]),
    ("2.6.4  Jurnal Umum",
     "Halaman Jurnal Umum diberi judul \u201cJurnal Umum\u201d (\U0001F4D2) dengan "
     "subjudul \u201cPencatatan kronologis seluruh transaksi sebelum penyesuaian "
     "(21 transaksi).\u201d Halaman ini menampilkan sebuah tabel besar dengan kolom "
     "Tanggal, Kode, Keterangan, Akun, Debit, dan Kredit, di mana setiap transaksi "
     "ditulis dalam baris terpisah untuk sisi debit dan sisi kreditnya. Di bagian bawah "
     "tabel terdapat baris Total Debit dan Total Kredit, dilengkapi badge (penanda) yang "
     "menunjukkan apakah kedua total sudah seimbang. Data yang ditampilkan adalah 21 "
     "transaksi mentah, yaitu gabungan data contoh bawaan (seed) dan transaksi yang "
     "diinput sendiri oleh pengguna, dan semuanya masih dalam kondisi belum disesuaikan. "
     "Dengan begitu, halaman ini menjadi catatan kronologis paling awal yang merekam apa "
     "adanya seluruh transaksi sesuai urutan terjadinya.",
     "04-jurnal-umum.png", "Tampilan Jurnal Umum.",
     ["Tabel memiliki enam kolom: Tanggal, Kode, Keterangan, Akun, Debit, dan Kredit.",
      "Setiap transaksi ditulis dalam baris terpisah untuk sisi debit (sisi penambahan) "
      "dan sisi kredit (sisi pengurangan/lawannya).",
      "Baris Total Debit dan Total Kredit di bawah tabel memastikan pencatatan berpasangan "
      "selalu seimbang.",
      "Badge keseimbangan menandai secara cepat apakah total debit sudah sama dengan "
      "total kredit.",
      "Isi tabel adalah 21 transaksi mentah (data contoh bawaan ditambah input pengguna) "
      "yang belum disesuaikan.",
      "Subjudul menyebutkan jumlah transaksi yang sedang ditampilkan sehingga mudah "
      "diketahui sekilas."]),
    ("2.6.5  Buku Besar",
     "Halaman Buku Besar berjudul \u201cBuku Besar\u201d (\U0001F4DA) dengan subjudul "
     "\u201cPengelompokan mutasi per akun beserta saldo berjalan.\u201d Berbeda dari "
     "jurnal umum yang berurut menurut waktu, halaman ini menyusun ulang transaksi "
     "berdasarkan akunnya. Setiap akun ditampilkan sebagai panel yang bisa dibuka-tutup "
     "(expander) dengan judul \u201c{kode} \u2014 {nama}\u201d, dan ketika dibuka akan "
     "menampilkan tabel berisi kolom Tanggal, Keterangan, Debit, Kredit, dan Saldo "
     "Berjalan, lalu ditutup baris Total Debit, Total Kredit, dan Saldo Akhir. Kolom "
     "Saldo Berjalan inilah inti halaman ini: nilainya merupakan akumulasi berjalan yang "
     "dihitung dari debit dikurangi kredit secara bertahap dari atas ke bawah, sehingga "
     "pengguna bisa melihat naik-turunnya saldo sebuah akun dari satu transaksi ke "
     "transaksi berikutnya. Dengan model panel yang bisa dilipat, daftar akun yang banyak "
     "tetap ringkas dan pengguna dapat fokus membuka hanya akun yang ingin ditelusuri.",
     "05-buku-besar.png", "Tampilan Buku Besar.",
     ["Setiap akun tampil sebagai panel lipat (expander) berjudul \u201c{kode} \u2014 "
      "{nama}\u201d yang bisa dibuka dan ditutup.",
      "Di dalam panel terdapat tabel dengan kolom Tanggal, Keterangan, Debit, Kredit, dan "
      "Saldo Berjalan.",
      "Saldo Berjalan adalah saldo yang terus diperbarui transaksi demi transaksi, "
      "dihitung dari debit dikurangi kredit secara berurutan.",
      "Bagian bawah panel menampilkan Total Debit, Total Kredit, dan Saldo Akhir per "
      "akun.",
      "Transaksi yang sebelumnya berurut waktu kini dikelompokkan rapi per akun agar "
      "mudah ditelusuri.",
      "Model panel lipat membuat daftar akun yang banyak tetap ringkas dan tidak memenuhi "
      "layar."]),
    ("2.6.6  Neraca Saldo",
     "Halaman Neraca Saldo berjudul \u201cNeraca Saldo\u201d (\u2696\uFE0F) dengan "
     "subjudul \u201cRingkasan saldo seluruh akun; total debit harus sama dengan total "
     "kredit.\u201d Halaman ini menyajikan satu tabel rangkuman dengan kolom Kode, Akun, "
     "Debit, dan Kredit, lalu ditutup dengan baris TOTAL di bagian bawah beserta badge "
     "(penanda) yang menunjukkan apakah kedua kolom sudah seimbang. Untuk menjaga "
     "tampilan tetap bersih dan mudah dibaca, akun yang saldonya nol otomatis "
     "disembunyikan sehingga hanya akun yang benar-benar bersaldo yang muncul. Angka pada "
     "tabel ini adalah saldo akhir tiap akun yang dirangkum dari seluruh transaksi, dan "
     "karena pencatatan dilakukan berpasangan, total debit dan total kredit selalu sama. "
     "Halaman ini berguna sebagai titik "
     "pemeriksaan untuk memastikan seluruh pembukuan masih dalam keadaan seimbang sebelum "
     "lanjut ke tahap penyesuaian.",
     "06-neraca-saldo.png", "Tampilan Neraca Saldo.",
     ["Tabel memiliki kolom Kode, Akun, Debit, dan Kredit yang merangkum saldo akhir tiap "
      "akun.",
      "Baris TOTAL di bawah tabel menjumlahkan seluruh debit dan seluruh kredit.",
      "Badge keseimbangan menegaskan total debit selalu sama dengan total kredit.",
      "Akun yang bersaldo nol disembunyikan otomatis agar tabel tetap ringkas dan mudah "
      "dibaca.",
      "Angka berasal dari rangkuman seluruh transaksi yang sudah dicatat sebelum "
      "penyesuaian.",
      "Halaman ini menjadi titik kontrol untuk memastikan pembukuan seimbang sebelum "
      "masuk tahap penyesuaian."]),
    ("2.6.7  Jurnal Penyesuaian",
     "Halaman Jurnal Penyesuaian berjudul \u201cJurnal Penyesuaian\u201d (\U0001F527) "
     "dengan subjudul \u201cPenyesuaian akhir periode agar saldo akun mencerminkan "
     "kondisi riil.\u201d Pada data contoh, halaman ini menampilkan satu entri "
     "penyesuaian bernama AJP1 tertanggal 30 April 2025, yaitu mencatat Debit pada akun "
     "Beban Perlengkapan sebesar Rp 100.000 dan Kredit pada akun Perlengkapan sebesar "
     "Rp 100.000. Di dekat tabel terdapat kotak catatan yang menjelaskan latar belakang "
     "penyesuaian tersebut, yakni perlengkapan berupa karung gabah senilai Rp 100.000 "
     "yang telah habis terpakai saat panen sehingga perlu dikurangi dari nilai aset dan "
     "diakui sebagai beban. Penyesuaian seperti ini diperlukan agar saldo akun benar-"
     "benar menggambarkan kondisi nyata di akhir periode, bukan sekadar angka pembelian "
     "di awal. Dengan begitu, laporan yang disusun setelah tahap ini menjadi lebih jujur "
     "dan akurat.",
     "07-jurnal-penyesuaian.png", "Tampilan Jurnal Penyesuaian.",
     ["Menampilkan entri penyesuaian AJP1 tertanggal 30 April 2025 dalam format tabel "
      "debit dan kredit.",
      "Isinya mencatat Debit Beban Perlengkapan Rp 100.000 dan Kredit Perlengkapan "
      "Rp 100.000.",
      "Tersedia kotak catatan yang menerangkan alasan penyesuaian dengan bahasa yang "
      "mudah dipahami.",
      "Latar belakangnya: perlengkapan karung gabah senilai Rp 100.000 telah habis "
      "terpakai saat panen.",
      "Penyesuaian membuat nilai aset berkurang dan memunculkan beban agar sesuai kondisi "
      "riil akhir periode.",
      "Hasil penyesuaian inilah yang akan dipakai oleh laporan-laporan berikutnya agar "
      "lebih akurat."]),
    ("2.6.8  Neraca Saldo Setelah Penyesuaian",
     "Halaman ini berjudul \u201cNeraca Saldo Setelah Penyesuaian\u201d (\u2696\uFE0F) "
     "dengan subjudul \u201cSaldo seluruh akun setelah jurnal penyesuaian diposting.\u201d "
     "Tampilannya mirip dengan Neraca Saldo biasa, yaitu berupa tabel saldo seluruh akun "
     "yang dilengkapi badge (penanda) keseimbangan, namun bedanya angka di sini sudah "
     "memperhitungkan jurnal penyesuaian yang dibuat sebelumnya. Artinya, jika di Jurnal "
     "Penyesuaian tadi tercatat pemakaian perlengkapan Rp 100.000, maka saldo akun "
     "Perlengkapan dan Beban Perlengkapan pada halaman ini sudah ikut berubah mengikuti "
     "penyesuaian tersebut. Angka pada tabel dihitung dari jurnal yang sudah disesuaikan, "
     "sehingga inilah dasar saldo yang lebih akurat untuk menyusun laporan keuangan "
     "utama seperti laba rugi, perubahan ekuitas, dan posisi keuangan. Halaman ini sekali "
     "lagi memastikan total debit tetap sama dengan total kredit setelah penyesuaian.",
     "08-ns-penyesuaian.png", "Tampilan Neraca Saldo Setelah Penyesuaian.",
     ["Menampilkan tabel saldo seluruh akun, serupa Neraca Saldo, namun angkanya sudah "
      "memperhitungkan penyesuaian.",
      "Saldo akun yang terdampak penyesuaian (misalnya Perlengkapan dan Beban "
      "Perlengkapan) sudah berubah mengikuti AJP1.",
      "Badge keseimbangan menegaskan total debit tetap sama dengan total kredit setelah "
      "penyesuaian.",
      "Angka dihitung otomatis dari jurnal yang sudah disesuaikan, bukan dari jurnal "
      "mentah.",
      "Saldo di halaman ini menjadi dasar penyusunan laba rugi, perubahan ekuitas, dan "
      "posisi keuangan.",
      "Halaman ini memastikan pembukuan tetap seimbang sebelum laporan keuangan utama "
      "disusun."]),
    ("2.6.9  Laporan Laba Rugi",
     "Halaman Laporan Laba Rugi berjudul \u201cLaporan Laba Rugi\u201d (\U0001F4C8) dan "
     "disusun menjadi dua seksi tabel utama. Seksi \u201cPendapatan\u201d memuat rincian "
     "sumber pemasukan beserta baris TOTAL pendapatannya, sedangkan seksi \u201cBeban\u201d "
     "memuat rincian seluruh biaya beserta baris TOTAL bebannya, masing-masing dengan "
     "kolom Akun dan Jumlah. Di bagian akhir terdapat kartu \u201cLaba Bersih\u201d yang "
     "berwarna hijau bila hasilnya untung (nol atau lebih) dan merah bila merugi, "
     "dilengkapi sub-baris yang memperlihatkan perhitungan \u201cPendapatan ... dikurangi "
     "Beban ...\u201d. Cara kerjanya sederhana: aplikasi menjumlahkan seluruh pendapatan, "
     "menjumlahkan seluruh beban, lalu mengurangkan keduanya, semuanya diambil dari "
     "jurnal yang sudah disesuaikan. Dengan penyajian dua seksi yang jelas plus kartu "
     "berwarna, pengguna langsung tahu apakah usaha taninya untung atau rugi pada periode "
     "tersebut.",
     "09-laba-rugi.png", "Tampilan Laporan Laba Rugi.",
     ["Seksi \u201cPendapatan\u201d menampilkan rincian pemasukan beserta baris TOTAL "
      "pendapatan.",
      "Seksi \u201cBeban\u201d menampilkan rincian biaya beserta baris TOTAL beban, "
      "dengan kolom Akun dan Jumlah.",
      "Kartu \u201cLaba Bersih\u201d berwarna hijau saat untung dan merah saat rugi "
      "sehingga mudah dikenali.",
      "Sub-baris di kartu memperlihatkan perhitungan Pendapatan dikurangi Beban secara "
      "ringkas.",
      "Angka dihitung otomatis dengan mengurangkan total beban dari total pendapatan pada "
      "jurnal yang sudah disesuaikan.",
      "Laporan ini menjawab pertanyaan utama pelaku usaha: untung atau rugi pada periode "
      "berjalan."]),
    ("2.6.10  Laporan Perubahan Ekuitas",
     "Halaman Laporan Perubahan Ekuitas berjudul \u201cLaporan Perubahan Ekuitas\u201d "
     "(\U0001F4B9) dengan subjudul \u201cPeriode panen 2025 \u2014 disusun dari jurnal "
     "yang telah disesuaikan.\u201d Halaman ini menampilkan tabel berkolom Keterangan dan "
     "Jumlah yang menelusuri perjalanan modal pemilik, mulai dari Modal Awal, lalu "
     "ditambah Laba Bersih, dikurangi Prive (pengambilan pribadi pemilik untuk keperluan "
     "di luar usaha), hingga akhirnya diperoleh Modal Akhir. Di bagian bawah disediakan "
     "kartu \u201cModal Akhir\u201d yang menonjolkan angka penutup tersebut. Perhitungan "
     "mengikuti rumus modal akhir = modal awal + laba bersih \u2212 prive, dengan nilai "
     "laba bersih diambil dari Laporan Laba Rugi dan seluruh angka bersumber dari jurnal "
     "yang telah disesuaikan. Dengan begitu, pengguna dapat melihat dengan jelas seberapa "
     "besar modalnya bertambah atau berkurang selama satu periode usaha tani.",
     "10-ekuitas.png", "Tampilan Laporan Perubahan Ekuitas.",
     ["Tabel berkolom Keterangan dan Jumlah menampilkan urutan Modal Awal, Laba Bersih "
      "(menambah), Prive (mengurangi), dan Modal Akhir.",
      "Prive adalah pengambilan dana pribadi oleh pemilik untuk keperluan di luar usaha, "
      "sehingga mengurangi modal.",
      "Tersedia kartu \u201cModal Akhir\u201d yang menonjolkan nilai modal di akhir "
      "periode.",
      "Perhitungan memakai rumus modal akhir = modal awal + laba bersih \u2212 prive.",
      "Nilai laba bersih diambil dari Laporan Laba Rugi dan seluruh angka bersumber dari "
      "jurnal yang telah disesuaikan.",
      "Laporan ini memperlihatkan secara jelas pertumbuhan atau penyusutan modal pemilik "
      "sepanjang periode."]),
    ("2.6.11  Laporan Posisi Keuangan (Neraca)",
     "Halaman Laporan Posisi Keuangan berjudul \u201cLaporan Posisi Keuangan "
     "(Neraca)\u201d (\U0001F3DB\uFE0F) dengan subjudul \u201cPer 30 April 2025 \u2014 "
     "Aset harus sama dengan Kewajiban + Ekuitas.\u201d Halaman ini ditata dalam dua "
     "kolom berdampingan: kolom kiri berjudul \u201cAset\u201d memuat rincian harta usaha "
     "beserta baris TOTAL ASET, sedangkan kolom kanan berjudul \u201cKewajiban & "
     "Ekuitas\u201d memuat rincian utang dan modal, termasuk Modal Akhir, beserta baris "
     "TOTAL-nya. Di antara kedua kolom terdapat badge (penanda) keseimbangan yang "
     "membandingkan total aset dengan total kewajiban ditambah ekuitas. Inti laporan ini "
     "adalah persamaan dasar akuntansi: apa yang dimiliki usaha (aset) harus sama dengan "
     "sumber pendanaannya (kewajiban ditambah ekuitas). Nilai Modal Akhir yang muncul di "
     "sini diambil dari Laporan Perubahan Ekuitas, sehingga seluruh laporan saling "
     "terhubung dan konsisten satu sama lain.",
     "11-posisi-keuangan.png", "Tampilan Laporan Posisi Keuangan.",
     ["Kolom kiri \u201cAset\u201d memuat rincian harta usaha beserta baris TOTAL ASET.",
      "Kolom kanan \u201cKewajiban & Ekuitas\u201d memuat rincian utang dan modal, "
      "termasuk Modal Akhir, beserta baris TOTAL.",
      "Badge keseimbangan membandingkan TOTAL ASET dengan total Kewajiban ditambah "
      "Ekuitas.",
      "Laporan menegaskan persamaan dasar akuntansi: Aset = Kewajiban + Ekuitas.",
      "Nilai Modal Akhir diambil dari Laporan Perubahan Ekuitas sehingga antarlaporan "
      "tetap konsisten.",
      "Subjudul menyebutkan tanggal pelaporan (per 30 April 2025) sebagai potret kondisi "
      "keuangan pada saat itu."]),
    ("2.6.12  Laporan Arus Kas",
     "Halaman Laporan Arus Kas berjudul \u201cLaporan Arus Kas\u201d (\U0001F4B5) dengan "
     "subjudul \u201cMetode langsung\u201d (cara penyusunan yang merinci penerimaan dan "
     "pengeluaran kas secara langsung). Halaman ini terbagi menjadi tiga seksi, yaitu "
     "Aktivitas Operasi (kas dari kegiatan usaha sehari-hari), Pendanaan (kas dari modal "
     "atau pinjaman), dan Investasi (kas untuk membeli atau menjual aset), di mana tiap "
     "seksi menampilkan rincian beserta totalnya. Di bagian bawah terdapat tabel "
     "ringkasan berisi Kenaikan (Penurunan) Kas Bersih, Kas Awal Periode (bernilai 0), "
     "dan Kas Akhir Periode, dilengkapi kartu \u201cKas Akhir Periode\u201d yang "
     "menonjolkan saldo kas penutup. Cara kerjanya: setiap transaksi yang menyentuh akun "
     "Kas ditelaah akun lawannya, lalu dikelompokkan ke salah satu dari tiga aktivitas "
     "tersebut, sehingga pengguna bisa melihat dari mana kas masuk dan ke mana kas keluar "
     "selama periode berjalan.",
     "12-arus-kas.png", "Tampilan Laporan Arus Kas.",
     ["Disusun dengan metode langsung, yaitu merinci penerimaan dan pengeluaran kas apa "
      "adanya.",
      "Terdapat tiga seksi: Aktivitas Operasi, Pendanaan, dan Investasi, masing-masing "
      "dengan rincian dan total.",
      "Tabel ringkasan memuat Kenaikan (Penurunan) Kas Bersih, Kas Awal Periode (0), dan "
      "Kas Akhir Periode.",
      "Tersedia kartu \u201cKas Akhir Periode\u201d yang menonjolkan saldo kas di akhir "
      "periode.",
      "Setiap transaksi yang menyentuh akun Kas dikelompokkan berdasarkan akun lawannya "
      "ke salah satu dari tiga aktivitas.",
      "Laporan ini memperlihatkan dengan jelas sumber masuknya kas dan tujuan keluarnya "
      "kas sepanjang periode."]),
    ("2.6.13  Jurnal Penutup",
     "Halaman Jurnal Penutup berjudul \u201cJurnal Penutup\u201d (\U0001F512) dengan "
     "subjudul \u201cMemindahkan saldo akun nominal ke modal agar bersaldo nol.\u201d "
     "Yang dimaksud akun nominal adalah akun pendapatan dan beban yang sifatnya sementara "
     "dan harus dinolkan di akhir periode agar periode berikutnya dimulai dari angka "
     "bersih. Halaman ini menampilkan tabel berisi empat entri penutupan, yaitu JP1 "
     "sampai JP4, yang angkanya dihitung otomatis oleh aplikasi. Di sampingnya disajikan "
     "daftar \u201cEmpat langkah penutupan\u201d sebagai panduan: JP1 menutup akun "
     "pendapatan ke Ikhtisar Laba Rugi, JP2 menutup akun beban ke Ikhtisar, JP3 "
     "memindahkan saldo Ikhtisar ke Modal, dan JP4 memindahkan Prive ke Modal. Pada "
     "narasi panduan disebutkan contoh angka laba Rp 11.500.000 dan prive Rp 4.500.000 "
     "sebagai ilustrasi agar mudah dibayangkan; angka ilustrasi ini sengaja ditulis tetap "
     "pada teks penjelasan, sedangkan yang benar-benar dihitung otomatis oleh aplikasi "
     "adalah tabel entri JP1 sampai JP4 tersebut.",
     "13-jurnal-penutup.png", "Tampilan Jurnal Penutup.",
     ["Tabel menampilkan empat entri penutupan JP1 sampai JP4 yang nilainya dihitung "
      "otomatis.",
      "Akun nominal (pendapatan dan beban) bersifat sementara dan dinolkan agar periode "
      "baru mulai dari angka bersih.",
      "Daftar \u201cEmpat langkah penutupan\u201d menjelaskan urutannya: JP1 tutup "
      "pendapatan ke Ikhtisar, JP2 tutup beban ke Ikhtisar.",
      "Lanjutannya: JP3 memindahkan saldo Ikhtisar ke Modal, dan JP4 memindahkan Prive "
      "ke Modal.",
      "Contoh angka laba Rp 11.500.000 dan prive Rp 4.500.000 hanya ilustrasi tetap pada "
      "narasi, bukan hasil hitung di halaman ini.",
      "Yang dihitung otomatis oleh aplikasi adalah tabel entri JP1 sampai JP4, bukan "
      "angka contoh pada narasi."]),
    ("2.6.14  Neraca Saldo Setelah Penutupan",
     "Halaman ini berjudul \u201cNeraca Saldo Setelah Penutupan\u201d (\u2696\uFE0F) "
     "dengan subjudul \u201cHanya akun riil yang tersisa; akun nominal sudah bersaldo "
     "nol.\u201d Setelah jurnal penutup dijalankan, akun-akun pendapatan dan beban sudah "
     "dinolkan, sehingga tabel pada halaman ini hanya menyaring dan menampilkan akun riil "
     "saja, yaitu akun aset, kewajiban, dan ekuitas, lengkap dengan badge (penanda) "
     "keseimbangan. Keterangan di bawahnya menjelaskan bahwa akun Pendapatan dan Beban "
     "kini bersaldo nol karena sudah dipindahkan ke modal melalui proses penutupan. "
     "Halaman ini menjadi bukti akhir bahwa siklus akuntansi satu periode telah tuntas "
     "dengan rapi: yang tersisa hanyalah akun-akun yang saldonya akan dibawa ke periode "
     "berikutnya, dan total debit tetap sama dengan total kredit. Dengan demikian, "
     "pembukuan siap memulai periode baru dari titik yang bersih.",
     "14-ns-penutupan.png", "Tampilan Neraca Saldo Setelah Penutupan.",
     ["Tabel hanya menampilkan akun riil, yaitu akun aset, kewajiban, dan ekuitas.",
      "Akun nominal (Pendapatan dan Beban) sudah bersaldo nol karena dipindahkan ke modal "
      "lewat jurnal penutup.",
      "Badge keseimbangan menegaskan total debit tetap sama dengan total kredit setelah "
      "penutupan.",
      "Keterangan di bawah tabel menjelaskan alasan akun Pendapatan dan Beban kini "
      "bernilai nol.",
      "Saldo yang tersisa adalah saldo yang akan dibawa sebagai modal awal periode "
      "berikutnya.",
      "Halaman ini menandai bahwa satu siklus akuntansi telah selesai dengan rapi dan "
      "seimbang."]),
    ("2.6.15  Input Transaksi",
     "Halaman Input Transaksi berjudul \u201c\u270F\uFE0F Input Transaksi\u201d dan diberi "
     "keterangan yang mengingatkan bahwa total debit harus sama dengan total kredit. "
     "Inilah satu-satunya halaman yang dapat menulis dan mengubah data. Di bagian atas "
     "terdapat form input yang otomatis kosong kembali setiap selesai menyimpan, berisi "
     "kolom Tanggal, kolom Keterangan (dengan contoh tulisan \u201cContoh: Pembelian "
     "benih padi\u201d), baris Debit (pilihan Akun Debit dari 15 akun resmi ditambah "
     "Nominal Debit yang naik turun per kelipatan 1000), baris Kredit (Akun Kredit "
     "ditambah Nominal Kredit), serta tombol \u201cSimpan Transaksi\u201d. Bila debit dan "
     "kredit tidak sama, penyimpanan ditolak disertai pesan kesalahan; bila sudah "
     "seimbang, muncul pesan \u201cTransaksi berhasil disimpan.\u201d dan aplikasi "
     "memberi kode transaksi otomatis berformat T diikuti angka. Di bawah form terdapat "
     "\u201c\U0001F4CB Daftar Transaksi\u201d berupa tabel Tanggal, Kode, Keterangan, "
     "Akun, Debit, dan Kredit, lalu \u201c\U0001F5C2\uFE0F Kelola Transaksi\u201d di mana "
     "setiap transaksi memiliki tombol \u201c\u270F\uFE0F Edit\u201d dan "
     "\u201c\U0001F5D1\uFE0F Hapus\u201d (semua transaksi, termasuk data contoh, bisa "
     "diedit atau dihapus). Tersedia pula panel \u201c\u26A0\uFE0F Reset ke Data Awal\u201d "
     "yang memuat peringatan, kotak centang konfirmasi, dan tombol \u201cReset "
     "Sekarang\u201d untuk mengembalikan 20 transaksi contoh.",
     "15-input-transaksi.png", "Halaman input transaksi baru.",
     ["Form input berisi kolom Tanggal, Keterangan (dengan contoh \u201cPembelian benih "
      "padi\u201d), baris Debit, dan baris Kredit.",
      "Akun Debit dan Akun Kredit dipilih dari 15 akun resmi, dan nominalnya naik turun "
      "per kelipatan 1000.",
      "Tombol \u201cSimpan Transaksi\u201d menolak data bila debit tidak sama dengan "
      "kredit, dan menerima bila seimbang dengan pesan \u201cTransaksi berhasil "
      "disimpan.\u201d",
      "Setiap transaksi tersimpan otomatis diberi kode berformat T diikuti angka, dan "
      "form langsung kosong kembali setelah menyimpan.",
      "Bagian \u201c\U0001F4CB Daftar Transaksi\u201d menampilkan tabel seluruh transaksi, "
      "sedangkan \u201c\U0001F5C2\uFE0F Kelola Transaksi\u201d menyediakan tombol Edit dan "
      "Hapus untuk tiap transaksi (termasuk data contoh).",
      "Panel \u201c\u26A0\uFE0F Reset ke Data Awal\u201d (dengan peringatan, centang "
      "konfirmasi, dan tombol \u201cReset Sekarang\u201d) mengembalikan 20 transaksi "
      "contoh bila diperlukan."]),
    ("2.6.16  Ubah Transaksi",
     "Ketika tombol \u201c\u270F\uFE0F Edit\u201d pada sebuah transaksi ditekan, aplikasi "
     "memunculkan popup mengambang (jendela dialog) di tengah layar berjudul "
     "\u201c\u270F\uFE0F Edit Transaksi {kode}\u201d. Form di dalam popup sudah otomatis "
     "terisi dengan data lama transaksi tersebut, meliputi Tanggal, Keterangan, Akun dan "
     "Nominal Debit, serta Akun dan Nominal Kredit, dilengkapi tombol \u201c\U0001F4BE "
     "Simpan Perubahan\u201d dan tombol \u201cBatal\u201d. Bila pengguna menyimpan dengan "
     "kondisi debit sama dengan kredit, perubahan tersimpan dan popup otomatis tertutup; "
     "namun bila debit tidak sama dengan kredit, pesan kesalahan tampil di dalam popup "
     "dan popup tetap terbuka tanpa mengubah data lama. Menekan \u201cBatal\u201d akan "
     "menutup popup tanpa perubahan apa pun. Popup ini memanfaatkan fitur dialog "
     "Streamlit sehingga selalu muncul mengambang di tengah layar terlepas dari posisi "
     "gulir, sebuah perbaikan dari versi sebelumnya yang formnya kerap muncul jauh dari "
     "tombol Edit dan menyulitkan pengguna.",
     "16-edit-transaksi.png", "Formulir ubah transaksi.",
     ["Tombol \u201c\u270F\uFE0F Edit\u201d memunculkan popup mengambang berjudul "
      "\u201c\u270F\uFE0F Edit Transaksi {kode}\u201d di tengah layar.",
      "Form di dalam popup sudah terisi data lama: Tanggal, Keterangan, Akun dan Nominal "
      "Debit, serta Akun dan Nominal Kredit.",
      "Tombol \u201c\U0001F4BE Simpan Perubahan\u201d menyimpan saat debit sama dengan "
      "kredit, lalu popup otomatis tertutup.",
      "Bila debit tidak sama dengan kredit, pesan kesalahan muncul di dalam popup dan "
      "popup tetap terbuka tanpa mengubah data lama.",
      "Tombol \u201cBatal\u201d menutup popup tanpa menyimpan perubahan apa pun.",
      "Popup memakai fitur dialog Streamlit agar selalu mengambang di tengah layar, "
      "memperbaiki versi lama yang formnya muncul jauh dari tombol Edit."]),
]

for title, intro, img, cap, bullets in shot_sections:
    h3(doc, title)
    para(doc, intro)
    figure(doc, SHOT / img, cap, width_cm=15.0)
    para(doc, "Penjelasan:", bold=True, justify=False)
    for b in bullets:
        bullet(doc, b)
    para(doc, "Dengan memahami isi dan cara kerja halaman ini, pengguna dapat membaca "
              "informasinya dengan percaya diri dan memanfaatkannya untuk mengelola "
              "keuangan usaha tani padi secara lebih tertib.")

page_break(doc)

h1(doc, "2.13  Fitur Stok / Persediaan")
para(doc, "Selain pencatatan keuangan berbasis jurnal, aplikasi Akuntansi Tani Padi "
          "juga menyediakan modul Stok / Persediaan yang membantu petani mengetahui "
          "sisa bahan usaha tani — benih, pupuk, pestisida, dan karung — beserta "
          "nilai rupiahnya. Fitur ini berdiri sendiri dari siklus akuntansi: stok "
          "tidak ditulis sebagai akun aset pada laporan keuangan, melainkan "
          "ditampilkan sebagai informasi persediaan yang dihitung ulang dari "
          "seluruh pergerakan masuk/keluar dengan metode harga rata-rata "
          "bergerak (moving average).")
para(doc, "Tujuan utama modul ini adalah:")
bullet(doc, "Membantu petani memantau ketersediaan bahan secara real-time sehingga "
            "kekurangan dapat diantisipasi sebelum masa tanam.")
bullet(doc, "Menyajikan nilai persediaan dalam rupiah sebagai gambaran modal "
            "yang tertahan di gudang.")
bullet(doc, "Menolak transaksi yang tidak logis, seperti pengeluaran melebihi sisa "
            "stok atau perubahan riwayat yang membuat saldo menjadi negatif.")
bullet(doc, "Menghubungkan pergerakan stok secara opsional ke jurnal umum melalui "
            "kolom referensi, tanpa mengganggu alur laporan keuangan.")

h2(doc, "2.13.1  Item Stok & Satuan")
para(doc, "Aplikasi menyediakan empat jenis item yang umum digunakan dalam "
          "usaha tani padi, masing-masing dengan satuan dan ambang stok minimum "
          "yang dapat disesuaikan. Satuan terbagi dua kelompok: nilai desimal "
          "(kilogram, liter) untuk bahan yang mungkin diukur pecahan, dan nilai "
          "bulat (lembar) untuk barang yang dihitung per unit.")
kv_table(doc, [
    ("Benih", "kg (desimal)"),
    ("Pupuk", "kg (desimal)"),
    ("Pestisida", "liter (desimal)"),
    ("Karung", "lembar (bulat)"),
], headers=("Item", "Satuan & Format"))

h2(doc, "2.13.2  Halaman Stok / Persediaan")
para(doc, "Halaman Stok / Persediaan dapat diakses dari sidebar dengan memilih "
          "menu \"📦 Stok/Persediaan\". Halaman ini menampilkan peringatan stok "
          "menipis di bagian atas, tabel ringkasan stok saat ini, formulir "
          "pencatatan pergerakan masuk/keluar, dan riwayat lengkap per item yang "
          "dapat dibuka sebagai expander. Setiap baris riwayat dilengkapi tombol "
          "Edit (muncul sebagai jendela dialog) dan Hapus.")
figure(doc, SHOT / "17-stok-halaman.png",
       "Tampilan Halaman Stok / Persediaan dengan peringatan stok menipis, "
       "tabel ringkasan, formulir pergerakan, dan expander riwayat.",
       width_cm=15.5)
para(doc, "Penjelasan:", bold=True, justify=False)
bullet(doc, "Peringatan kuning di atas tabel otomatis muncul bila ada item yang "
            "sisa kuantitasnya kurang dari atau sama dengan ambang stok minimum.")
bullet(doc, "Tabel ringkasan menunjukkan sisa, nilai rupiah, dan status setiap "
            "item; nilai dihitung ulang dari pergerakan, bukan disimpan.")
bullet(doc, "Formulir \"Catat Pergerakan\" menerima tanggal, tipe (masuk/keluar), "
            "kuantitas, dan harga satuan. Untuk tipe \"keluar\" aplikasi secara "
            "otomatis menentukan nilai keluar dari harga rata-rata bergerak "
            "tanggal tersebut, sehingga kolom harga satuan diabaikan.")
bullet(doc, "Riwayat per item memperlihatkan setiap pergerakan dengan saldo "
            "berjalan (qty, nilai) sehingga pengguna dapat menelusuri "
            "perubahan stok dari waktu ke waktu.")
bullet(doc, "Edit dilakukan pada jendela popup yang mempertahankan nilai lama "
            "sebagai nilai awal; Hapus akan menampilkan pesan kesalahan bila "
            "penghapusan menyebabkan saldo negatif di titik manapun.")

h2(doc, "2.13.3  Kartu Ringkasan Stok di Dashboard")
para(doc, "Halaman Dashboard turut menampilkan tiga kartu ringkasan stok di "
          "bawah ringkasan keuangan: jumlah jenis item, jumlah item yang "
          "menipis, dan total nilai persediaan dalam rupiah. Sesuai prinsip "
          "pemisahan modul, nilai stok TIDAK ikut serta dalam perhitungan "
          "laba rugi, posisi keuangan, maupun arus kas — nilai tersebut murni "
          "informasi persediaan.")
figure(doc, SHOT / "18-dashboard-stok.png",
       "Dashboard dengan kartu keuangan (atas) dan kartu ringkasan stok (bawah).",
       width_cm=15.5)
para(doc, "Catatan penting:", bold=True, justify=False)
bullet(doc, "Kartu \"Total Nilai Stok\" dihasilkan langsung dari koneksi "
            "database (db.get_stok_ringkasan) dan tidak melewati alur "
            "get_data() yang dipakai laporan keuangan.")
bullet(doc, "Tujuan pemisahan ini agar penambahan atau pengurangan stok tidak "
            "mengubah laba/rugi, melainkan hanya ditampilkan sebagai "
            "informasi modal yang tertahan di gudang.")

h2(doc, "2.13.4  Validasi Stok Negatif")
para(doc, "Modul stok menerapkan pola validasi-sebelum-tulis: setiap "
          "penambahan, perubahan, atau penghapusan pergerakan diperiksa "
          "terlebih dahulu menggunakan mesin replay seluruh riwayat. Bila "
          "di titik manapun saldo menjadi negatif, transaksi ditolak dengan "
          "pesan kesalahan yang jelas, dan database tetap utuh (tidak ada "
          "penulisan sebagian).")
figure(doc, SHOT / "19-stok-tolak-negatif.png",
       "Contoh penolakan transaksi keluar yang melebihi sisa stok; database "
       "tidak berubah karena validasi dilakukan sebelum penulisan.",
       width_cm=15.5)
para(doc, "Aturan validasi yang berlaku:", bold=True, justify=False)
bullet(doc, "Keluar dengan kuantitas melebihi sisa ditolak.")
bullet(doc, "Edit pergerakan yang membuat saldo negatif pada titik manapun "
            "di riwayat ditolak; nilai lama tetap tersimpan.")
bullet(doc, "Hapus pergerakan yang membuat saldo negatif pada titik manapun "
            "di riwayat ditolak.")
bullet(doc, "Pergerakan divalidasi secara kronologis menggunakan metode "
            "harga rata-rata bergerak, sehingga nilai keluar selalu "
            "konsisten dengan rata-rata historis saat itu.")

page_break(doc)

# ===== BAB III PENUTUP =====
h1(doc, "BAB III  PENUTUP")

h2(doc, "3.1  Kesimpulan")
para(doc, "Aplikasi Akuntansi Tani Padi berhasil menerapkan sistem informasi akuntansi "
          "yang lengkap untuk kebutuhan pencatatan keuangan usaha tani padi. Aplikasi "
          "mampu mencatat transaksi dengan sistem berpasangan, memvalidasi keseimbangan "
          "debit dan kredit, serta menyajikan siklus akuntansi penuh dalam sebelas "
          "tahap laporan, mulai dari jurnal umum hingga neraca saldo setelah "
          "penutupan.")
para(doc, "Dengan antarmuka berbasis web yang sederhana dan tema yang konsisten, "
          "aplikasi ini memudahkan pengguna tanpa latar belakang akuntansi yang "
          "mendalam untuk mengelola keuangan usahanya. Penggunaan tipe Decimal dan "
          "penyimpanan nilai uang sebagai teks menjamin akurasi perhitungan, sementara "
          "pengamanan kata sandi dengan PBKDF2-HMAC-SHA256 memberikan perlindungan pada "
          "data akun pengguna.")

h2(doc, "3.2  Saran")
numbered(doc, "Menambahkan fitur ekspor laporan ke format PDF atau Excel agar laporan "
              "mudah dibagikan dan diarsipkan.")
numbered(doc, "Menyediakan fitur multi-periode sehingga pengguna dapat membandingkan "
              "kinerja keuangan antarmusim tanam.")
numbered(doc, "Menambahkan grafik visual pada dasbor untuk mempermudah pembacaan tren "
              "pendapatan dan beban.")
numbered(doc, "Mempertimbangkan penyimpanan berbasis daring (cloud) agar data dapat "
              "diakses dari berbagai perangkat secara aman.")

# ===== FOOTER =====
footer_with_pagenum(sec0, "Buku Panduan Aplikasi Akuntansi Tani Padi")

enable_update_fields(doc)
doc.save(str(OUT))

n_imgs = len(doc.inline_shapes)
print(f"OK  {OUT.name}")
print(f"    inline_shapes = {n_imgs}")
print(f"    figures (Gambar) = {_fig['n']}")
print(f"    size = {OUT.stat().st_size / 1024:.0f} KB")
